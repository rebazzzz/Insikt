from __future__ import annotations

import os
import re
import tempfile

from insikt.rag import extract_citations_from_response


def build_export_sections(text: str, lang: str = "en") -> dict:
    citations = extract_citations_from_response(text)
    citation_order = []
    citation_map = {}
    appendix_sources = []
    appendix_seen = set()

    def replacer(match):
        source = match.group(2).strip()
        page = match.group(4).strip()
        key = (source, page)
        if key not in citation_map:
            citation_map[key] = len(citation_order) + 1
            citation_order.append(key)
        if key not in appendix_seen:
            appendix_seen.add(key)
            appendix_sources.append(key)
        return f"[{citation_map[key]}]"

    body = re.sub(
        r"\[(Källa|Source):\s*([^,\]]+),\s*(sida|page)\s*([^\]]+)\]",
        replacer,
        text,
        flags=re.IGNORECASE,
    )

    endnotes_title = "Endnotes" if lang == "en" else "Slutnoter"
    appendix_title = "Source appendix" if lang == "en" else "Källbilaga"
    page_label = "page" if lang == "en" else "sida"

    endnotes = [f"[{index}] {source}, {page_label} {page}" for index, (source, page) in enumerate(citation_order, start=1)]
    appendix = [f"- {source}, {page_label} {page}" for source, page in appendix_sources]

    sections = [body.strip()]
    if endnotes:
        sections.append(f"{endnotes_title}\n" + "\n".join(endnotes))
    if appendix:
        sections.append(f"{appendix_title}\n" + "\n".join(appendix))

    return {
        "body": body.strip(),
        "endnotes_title": endnotes_title,
        "appendix_title": appendix_title,
        "endnotes": endnotes,
        "appendix": appendix,
        "full_text": "\n\n".join(section for section in sections if section).strip(),
    }


def export_text(text, lang: str = "en", safe_mode: bool = True):
    if not safe_mode:
        return text.encode("utf-8")
    return build_export_sections(text, lang)["full_text"].encode("utf-8")


def export_docx(text, lang: str = "en", safe_mode: bool = True):
    from docx import Document as DocxDocument

    doc = DocxDocument()
    if not safe_mode:
        for paragraph in text.split("\n"):
            doc.add_paragraph(paragraph)
    else:
        package = build_export_sections(text, lang)
        for paragraph in package["body"].split("\n"):
            doc.add_paragraph(paragraph)
        if package["endnotes"]:
            doc.add_heading(package["endnotes_title"], level=1)
            for note in package["endnotes"]:
                doc.add_paragraph(note)
        if package["appendix"]:
            doc.add_heading(package["appendix_title"], level=1)
            for item in package["appendix"]:
                doc.add_paragraph(item)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        with open(tmp.name, "rb") as handle:
            data = handle.read()
    os.unlink(tmp.name)
    return data


def export_pdf(text, lang: str = "en", safe_mode: bool = True):
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    output_text = build_export_sections(text, lang)["full_text"] if safe_mode else text
    for line in output_text.split("\n"):
        pdf.multi_cell(0, 10, line)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        pdf.output(tmp.name)
        with open(tmp.name, "rb") as handle:
            data = handle.read()
    os.unlink(tmp.name)
    return data


def export_markdown(text, lang: str = "en", safe_mode: bool = True):
    if not safe_mode:
        return text.encode("utf-8")
    package = build_export_sections(text, lang)
    markdown = package["body"]
    if package["endnotes"]:
        markdown += f"\n\n## {package['endnotes_title']}\n" + "\n".join(f"- {note}" for note in package["endnotes"])
    if package["appendix"]:
        markdown += f"\n\n## {package['appendix_title']}\n" + "\n".join(package["appendix"])
    return markdown.encode("utf-8")
