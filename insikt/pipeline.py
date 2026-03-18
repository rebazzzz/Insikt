from __future__ import annotations

import hashlib
import importlib.util
import os
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from shutil import which
from typing import Callable, List, Optional, Sequence, Tuple

from docx import Document as DocxDocument
from langchain_core.documents import Document

from .common import compute_uploaded_files_fingerprint, docs_to_records, read_json, records_to_docs, write_json


StatusCallback = Optional[Callable[[str], None]]
ErrorCallback = Optional[Callable[[str], None]]
ProgressCallback = Optional[Callable[[float], None]]


def _uploaded_file_cache_key(uploaded_file) -> str:
    hasher = hashlib.md5()
    content = uploaded_file.getvalue()
    hasher.update(uploaded_file.name.encode("utf-8", errors="ignore"))
    hasher.update(str(len(content)).encode("ascii"))
    hasher.update(content)
    return hasher.hexdigest()


def ocr_stack_available() -> bool:
    return (
        importlib.util.find_spec("pytesseract") is not None
        and importlib.util.find_spec("pypdfium2") is not None
        and which("tesseract") is not None
    )


def pdf_needs_ocr(pages: Sequence[Document]) -> bool:
    if not pages:
        return True
    meaningful_pages = 0
    for page in pages[:5]:
        text = " ".join((page.page_content or "").split())
        if len(text) >= 40:
            meaningful_pages += 1
    return meaningful_pages == 0


def ocr_pdf_file(pdf_path: str, source_name: str) -> List[Document]:
    import pypdfium2 as pdfium
    import pytesseract

    pdf = pdfium.PdfDocument(pdf_path)
    docs = []
    for page_index in range(len(pdf)):
        page = pdf[page_index]
        bitmap = page.render(scale=2)
        image = bitmap.to_pil()
        text = pytesseract.image_to_string(image).strip()
        if text:
            docs.append(Document(page_content=text, metadata={"source": source_name, "page": str(page_index + 1), "ocr": True}))
    return docs


def _load_single_file_cached(uploaded_file, cache_root: Path, error_callback: ErrorCallback = None, status_callback: StatusCallback = None) -> tuple[list[Document], bool]:
    cache_key = _uploaded_file_cache_key(uploaded_file)
    cache_file = cache_root / "file_pages" / cache_key / "raw_pages.json"
    if cache_file.exists():
        return records_to_docs(read_json(cache_file, [])), True
    docs = load_single_file(uploaded_file, error_callback=error_callback, status_callback=status_callback)
    write_json(cache_file, docs_to_records(docs))
    return docs, False


def get_cache_stats(cache_root: Path) -> dict:
    cache_root.mkdir(parents=True, exist_ok=True)
    bundle_cache_count = 0
    file_cache_count = 0
    for item in cache_root.iterdir():
        if not item.is_dir():
            continue
        if item.name == "file_pages":
            file_cache_count = sum(1 for child in item.iterdir() if child.is_dir())
        else:
            bundle_cache_count += 1
    return {
        "bundle_caches": bundle_cache_count,
        "file_caches": file_cache_count,
    }


def load_single_pdf(uploaded_file, error_callback: ErrorCallback = None, status_callback: StatusCallback = None) -> List[Document]:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(uploaded_file.getvalue())
        tmp_path = tmp.name
    try:
        from langchain_community.document_loaders import PyPDFLoader

        loader = PyPDFLoader(tmp_path)
        pages = loader.load()
        for page in pages:
            page.metadata["source"] = uploaded_file.name
            page.metadata.setdefault("page", page.metadata.get("page", "?"))
        if pdf_needs_ocr(pages) and ocr_stack_available():
            if status_callback:
                status_callback(f"Running OCR for scanned PDF: {uploaded_file.name}")
            ocr_pages = ocr_pdf_file(tmp_path, uploaded_file.name)
            if ocr_pages:
                return ocr_pages
        return pages
    except Exception:
        if error_callback:
            error_callback(uploaded_file.name)
        return []
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def load_single_text_file(uploaded_file) -> List[Document]:
    raw = uploaded_file.getvalue()
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("latin-1", errors="ignore")
    return [Document(page_content=text, metadata={"source": uploaded_file.name, "page": "1"})]


def load_single_docx(uploaded_file) -> List[Document]:
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(uploaded_file.getvalue())
            tmp_path = tmp.name
        doc = DocxDocument(tmp_path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return [Document(page_content=text, metadata={"source": uploaded_file.name, "page": "1"})]
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


def load_single_file(uploaded_file, error_callback: ErrorCallback = None, status_callback: StatusCallback = None) -> List[Document]:
    suffix = Path(uploaded_file.name).suffix.lower()
    if suffix == ".pdf":
        return load_single_pdf(uploaded_file, error_callback=error_callback, status_callback=status_callback)
    if suffix in [".txt", ".md"]:
        return load_single_text_file(uploaded_file)
    if suffix == ".docx":
        return load_single_docx(uploaded_file)
    return []


def semantic_chunking(
    pages: Sequence[Document],
    embeddings,
    chunk_size: int,
    chunk_overlap: int,
    status_callback: StatusCallback = None,
) -> List[Document]:
    import numpy as np
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    if status_callback:
        status_callback("Creating semantic segments...")

    sentence_splitter = RecursiveCharacterTextSplitter(
        chunk_size=300,
        chunk_overlap=50,
        separators=[". ", "! ", "? ", "\n"],
    )

    all_sentences = []
    for page in pages:
        for sentence in sentence_splitter.split_text(page.page_content):
            if sentence.strip():
                all_sentences.append(
                    {
                        "text": sentence,
                        "source": page.metadata.get("source", "Unknown"),
                        "page": page.metadata.get("page", "?"),
                    }
                )

    if not all_sentences:
        splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        return splitter.split_documents(list(pages))

    if status_callback:
        status_callback("Calculating semantic similarity...")

    embeddings_matrix = embeddings.embed_documents([item["text"] for item in all_sentences])
    chunks = []
    current_texts = []
    current_sources = set()
    current_pages = set()

    for index, sentence in enumerate(all_sentences):
        current_texts.append(sentence["text"])
        current_sources.add(sentence["source"])
        current_pages.add(str(sentence["page"]))

        similarity = 1.0
        if index > 0:
            prev_embed = embeddings_matrix[index - 1]
            curr_embed = embeddings_matrix[index]
            prev_norm = np.linalg.norm(prev_embed)
            curr_norm = np.linalg.norm(curr_embed)
            similarity = float(np.dot(prev_embed, curr_embed) / (prev_norm * curr_norm)) if prev_norm and curr_norm else 0.0

        if similarity < 0.5 and current_texts:
            chunks.append(
                Document(
                    page_content=" ".join(current_texts),
                    metadata={"source": ", ".join(sorted(current_sources)), "page": ", ".join(sorted(current_pages))},
                )
            )
            current_texts = []
            current_sources = set()
            current_pages = set()

    if current_texts:
        chunks.append(
            Document(
                page_content=" ".join(current_texts),
                metadata={"source": ", ".join(sorted(current_sources)), "page": ", ".join(sorted(current_pages))},
            )
        )

    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    normalized = []
    for chunk in chunks:
        if len(chunk.page_content) > chunk_size * 2:
            for piece in splitter.split_text(chunk.page_content):
                normalized.append(Document(page_content=piece, metadata=chunk.metadata))
        else:
            normalized.append(chunk)
    return normalized


def rechunk_pages(
    pages: Sequence[Document],
    embeddings,
    chunking_strategy: str,
    chunk_size: int,
    chunk_overlap: int,
    status_callback: StatusCallback = None,
) -> List[Document]:
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    if chunking_strategy == "semantic":
        return semantic_chunking(pages, embeddings, chunk_size, chunk_overlap, status_callback=status_callback)
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return splitter.split_documents(list(pages))


def process_uploaded_files(
    uploaded_files: Sequence,
    embeddings,
    cache_root: Path,
    chunking_strategy: str,
    chunk_size: int,
    chunk_overlap: int,
    status_callback: StatusCallback = None,
    progress_callback: ProgressCallback = None,
    error_callback: ErrorCallback = None,
    use_file_cache: bool = True,
) -> Tuple[str, List[Document], List[Document], dict]:
    fingerprint = compute_uploaded_files_fingerprint(uploaded_files)
    cache_dir = cache_root / fingerprint
    raw_cache = cache_dir / "raw_pages.json"
    chunk_cache = cache_dir / "chunks.json"

    if raw_cache.exists() and chunk_cache.exists():
        return fingerprint, records_to_docs(read_json(raw_cache, [])), records_to_docs(read_json(chunk_cache, [])), {
            "bundle_cache_hit": True,
            "file_cache_hits": len(uploaded_files),
            "files_processed": 0,
        }

    raw_pages: List[Document] = []
    stats = {"bundle_cache_hit": False, "file_cache_hits": 0, "files_processed": 0}
    total_files = max(len(uploaded_files), 1)
    with ThreadPoolExecutor(max_workers=min(4, total_files)) as executor:
        if use_file_cache:
            futures = {
                executor.submit(_load_single_file_cached, uploaded_file, cache_root, error_callback, status_callback): uploaded_file
                for uploaded_file in uploaded_files
            }
        else:
            futures = {executor.submit(load_single_file, uploaded_file, error_callback, status_callback): uploaded_file for uploaded_file in uploaded_files}
        for index, future in enumerate(as_completed(futures), start=1):
            result = future.result()
            if use_file_cache:
                pages, cache_hit = result
                if cache_hit:
                    stats["file_cache_hits"] += 1
                else:
                    stats["files_processed"] += 1
            else:
                pages = result
                stats["files_processed"] += 1
            raw_pages.extend(pages)
            if progress_callback:
                progress_callback(index / total_files)
            if status_callback:
                status_callback(f"Reading documents ({index}/{total_files})")

    chunks = rechunk_pages(
        raw_pages,
        embeddings,
        chunking_strategy=chunking_strategy,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        status_callback=status_callback,
    )

    write_json(raw_cache, docs_to_records(raw_pages))
    write_json(chunk_cache, docs_to_records(chunks))
    return fingerprint, raw_pages, chunks, stats


def build_or_load_vectorstore(
    fingerprint: str,
    chunks: Sequence[Document],
    embeddings,
    cache_root: Path,
    status_callback: StatusCallback = None,
) -> object:
    cache_dir = cache_root / fingerprint / "vectorstore"
    from langchain_community.vectorstores import FAISS

    if (cache_dir / "index.faiss").exists() and (cache_dir / "index.pkl").exists():
        return FAISS.load_local(str(cache_dir), embeddings, allow_dangerous_deserialization=True)

    if status_callback:
        status_callback("Building knowledge base...")
    texts = [doc.page_content for doc in chunks]
    metadatas = [doc.metadata for doc in chunks]
    vectorstore = FAISS.from_texts(texts, embeddings, metadatas=metadatas)
    cache_dir.mkdir(parents=True, exist_ok=True)
    vectorstore.save_local(str(cache_dir))
    return vectorstore
