"""
Insikt – Journalist AI Suite
GPU-accelerated, 100% local, free. Swedish/English bilingual.
Features: Background summarization, real-time chat with document RAG, entity extraction, timeline, sentiment, export.
Now with user-selectable GPU/CPU processing.
"""

import os
import subprocess
import sys
import tempfile
import pickle
import re
import shutil
import threading
import time
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional
from urllib.parse import quote
import streamlit as st
from streamlit.runtime.uploaded_file_manager import UploadedFile
import streamlit.components.v1 as components

# GPU detection
import torch

# Parallel loading
from concurrent.futures import ThreadPoolExecutor, as_completed

# LangChain
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_ollama import ChatOllama
from langchain_core.prompts import PromptTemplate
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage, AIMessage

# Export
from docx import Document as DocxDocument
from fpdf import FPDF

# Analysis tools
from transformers import pipeline
import yake

from insikt.analysis import analyze_sentiment, bias_check, compare_sources, extract_claim_check_items, extract_entities, extract_keywords, extract_quote_candidates, extract_timeline, translate_text
from insikt import get_build_metadata
from insikt.common import cleaned_ui_text, docs_to_records, records_to_docs, safe_html_fragment
from insikt.exports import export_docx, export_markdown, export_pdf, export_text
from insikt.feedback_store import build_feedback_bundle, list_issue_reports, save_issue_report
from insikt.pipeline import build_or_load_vectorstore, get_cache_stats, process_uploaded_files as cached_process_uploaded_files, rechunk_pages
from insikt.rag import chat_with_docs as rag_chat_with_docs
from insikt.rag import generate_writing as rag_generate_writing
from insikt.rag import assess_answer_confidence
from insikt.session_store import delete_slot, list_save_slots, load_slot, save_slot
from insikt.summarization import SummaryThread as ModularSummaryThread
from insikt.summarization import generate_doc_hash as modular_generate_doc_hash
from insikt.ui import app_readiness_label, concise_model_label
from insikt.validation import (
    get_missing_python_packages,
    get_installed_ollama_models,
    get_model_recommendations,
    get_system_profile,
    get_tesseract_install_hint,
    install_missing_python_packages,
    resolve_installed_ollama_model,
    run_startup_checks,
)

# -------------------------------------------------------------------
# Configuration & Language
# -------------------------------------------------------------------
APP_NAME = "Insikt"
CACHE_ROOT = Path("session_data/cache")
SAVES_ROOT = Path("session_data/saves")
FEEDBACK_ROOT = Path("session_data/feedback")

# LLM Models - Optimized for performance and accuracy
# Quantized models are smaller, faster, and use less memory while maintaining good quality
LLM_MODELS = {
    "llama3.2": {
        "display_name": "Llama 3.2 (Standard)",
        "display_name_en": "Llama 3.2 (Standard)",
        "description": "Bästa kvalitet men kräver mer minne. Rekommenderas för kraftfulla datorer.",
        "description_en": "Best quality but requires more memory. Recommended for powerful computers.",
        "speed": "Hastighet: Medel",
        "quality": "Kvalitet: Mycket hög",
        "memory": "High",
    },
    "llama3.2:3b": {
        "display_name": "Llama 3.2 3B (Snabb)",
        "display_name_en": "Llama 3.2 3B (Fast)",
        "description": "Snabbare och mindre minneskrävande. Bra balans mellan hastighet och kvalitet.",
        "description_en": "Faster and less memory intensive. Good balance between speed and quality.",
        "speed": "Hastighet: Hög",
        "quality": "Kvalitet: Hög",
        "memory": "Medium",
    },
    "llama3.2:1b": {
        "display_name": "Llama 3.2 1B (Turbo)",
        "display_name_en": "Llama 3.2 1B (Turbo)",
        "description": "Snabbast och minst minneskrävande. För svagare datorer eller snabba tester.",
        "description_en": "Fastest and most memory efficient. For weaker computers or quick tests.",
        "speed": "Hastighet: Mycket hög",
        "quality": "Kvalitet: Medel",
        "memory": "Low",
    },
    "mistral:7b": {
        "display_name": "Mistral 7B",
        "display_name_en": "Mistral 7B",
        "description": "Alternativ modell med bra prestanda. Bra för engelska dokument.",
        "description_en": "Alternative model with good performance. Good for English documents.",
        "speed": "Hastighet: Medel",
        "quality": "Kvalitet: Hög",
        "memory": "High",
    },
}
DEFAULT_LLM_MODEL = "llama3.2:latest"

# Chunking strategies
CHUNKING_STRATEGIES = {
    "semantic": {
        "display_name": "Smart (Semantisk)",
        "display_name_en": "Smart (Semantic)",
        "description": "Delar dokument baserat på mening och sammanhang. Bäst för längre dokument.",
        "description_en": "Splits documents based on meaning and context. Best for longer documents.",
        "icon": "",
    },
    "fixed": {
        "display_name": "Standard (Fast storlek)",
        "display_name_en": "Standard (Fixed size)",
        "description": "Delar dokument i jämna delar. Snabbare men mindre flexibelt.",
        "description_en": "Splits documents into equal parts. Faster but less flexible.",
        "icon": "",
    },
}
DEFAULT_CHUNKING = "semantic"

OLLAMA_MODEL = DEFAULT_LLM_MODEL
CHUNK_SIZE = 1500
CHUNK_OVERLAP = 500

# Embedding models - bge provides better quality than MiniLM
# User-friendly explanations added for non-technical users
EMBEDDING_MODELS = {
    "bge-small": {
        "model_name": "BAAI/bge-small-en-v1.5",
        "display_name": "Snabb (bge-small)",  # Fast (bge-small)
        "display_name_en": "Fast (bge-small)",
        "description": "Snabbare, mindre noggrann. Bra för testning och svaga datorer.",  # Faster, less accurate. Good for testing and weak computers.
        "description_en": "Faster, less accurate. Good for testing and weaker computers.",
        "speed": "Hastighet: Mycket hög",
        "quality": "Kvalitet: Låg",
    },
    "bge-base": {
        "model_name": "BAAI/bge-base-en-v1.5",
        "display_name": "Balanserad (bge-base)",  # Balanced (bge-base)
        "display_name_en": "Balanced (bge-base)",
        "description": "Bästa valet! Balans mellan hastighet och noggrannhet.",  # Best choice! Balance between speed and accuracy.
        "description_en": "Best choice! Balance between speed and accuracy.",
        "speed": "Hastighet: Hög",
        "quality": "Kvalitet: Medel",
    }
}
DEFAULT_EMBEDDING_MODEL = "bge-base"

# Default device (will be overridden by user choice)
DEFAULT_DEVICE = "cuda" if torch.cuda.is_available() else "cpu"

# Writing Studio options
WRITING_ROLES = {
    "author": {"sv": "Forfattare", "en": "Author"},
    "dramaturg": {"sv": "Dramaturg", "en": "Dramaturg"},
    "writer": {"sv": "Skribent", "en": "Writer"},
    "editor": {"sv": "Redaktor", "en": "Editor"},
}
WRITING_FORMATS = {
    "book": {"sv": "Bokmanus", "en": "Book Script"},
    "documentary": {"sv": "Dokumentarfilm", "en": "Documentary"},
    "screenplay": {"sv": "Manus (scener)", "en": "Screenplay"},
    "article": {"sv": "Artikel", "en": "Article"},
}
WRITING_TONES = {
    "neutral": {"sv": "Neutral", "en": "Neutral"},
    "investigative": {"sv": "Granskande", "en": "Investigative"},
    "narrative": {"sv": "Berättande", "en": "Narrative"},
    "formal": {"sv": "Formell", "en": "Formal"},
    "literary": {"sv": "Litterär", "en": "Literary"},
    "poetic": {"sv": "Poetisk", "en": "Poetic"},
    "immersive": {"sv": "Närvarande", "en": "Immersive"},
}
WRITING_LENGTHS = {
    "short": {"sv": "Kort", "en": "Short", "words": 400},
    "medium": {"sv": "Medel", "en": "Medium", "words": 800},
    "long": {"sv": "Lång", "en": "Long", "words": 1500},
}

REPORTER_TEMPLATES = {
    "summarize_interview": {
        "title": {"sv": "Sammanfatta intervju", "en": "Summarize interview"},
        "summary_focus": {
            "sv": "Sammanfatta intervjun tydligt. Lyft fram huvudpåståenden, viktiga citat, vem som säger vad, osäkerheter och konkreta uppföljningsfrågor.",
            "en": "Summarize the interview clearly. Highlight the main claims, important quotes, who says what, uncertainties, and concrete follow-up questions.",
        },
        "summary_style": "neutral",
        "summary_target_pages": 4,
        "summary_words_per_page": 260,
        "writing_brief": {
            "sv": "Skriv en kort intervjusammanfattning med huvudpunkter, citat att dubbelkolla och möjliga vinklar för uppföljning.",
            "en": "Write a short interview summary with key points, quotes to verify, and possible follow-up angles.",
        },
        "writing_format": "article",
        "writing_tone": "neutral",
        "writing_length": "medium",
    },
    "build_timeline": {
        "title": {"sv": "Bygg tidslinje", "en": "Build timeline"},
        "summary_focus": {
            "sv": "Ordna materialet kronologiskt. Fokusera på datum, händelser, aktörer, vad som förändras över tid och luckor i tidslinjen.",
            "en": "Arrange the material chronologically. Focus on dates, events, actors, what changes over time, and gaps in the timeline.",
        },
        "summary_style": "formal",
        "summary_target_pages": 3,
        "summary_words_per_page": 240,
        "writing_brief": {
            "sv": "Bygg en kronologisk tidslinje med datum, händelse, källa och vad som fortfarande behöver bekräftas.",
            "en": "Build a chronological timeline with date, event, source, and what still needs confirmation.",
        },
        "writing_format": "article",
        "writing_tone": "formal",
        "writing_length": "medium",
    },
    "find_contradictions": {
        "title": {"sv": "Hitta motsägelser", "en": "Find contradictions"},
        "summary_focus": {
            "sv": "Leta efter motsägelser mellan dokument, olika versioner av samma händelse, avvikande siffror och påståenden som inte går ihop.",
            "en": "Look for contradictions between documents, conflicting versions of the same event, inconsistent figures, and claims that do not line up.",
        },
        "summary_style": "investigative",
        "summary_target_pages": 4,
        "summary_words_per_page": 260,
        "writing_brief": {
            "sv": "Lista dokumenterade motsägelser, vilka källor som krockar och vilka frågor som bör ställas för att reda ut skillnaderna.",
            "en": "List documented contradictions, which sources conflict, and which questions should be asked to resolve the differences.",
        },
        "writing_format": "article",
        "writing_tone": "investigative",
        "writing_length": "medium",
    },
    "extract_names_roles": {
        "title": {"sv": "Namn och roller", "en": "Extract names and roles"},
        "summary_focus": {
            "sv": "Identifiera personer, organisationer, titlar, ansvar och relationer mellan aktörerna i materialet.",
            "en": "Identify people, organizations, titles, responsibilities, and relationships between the actors in the material.",
        },
        "summary_style": "neutral",
        "summary_target_pages": 2,
        "summary_words_per_page": 220,
        "writing_brief": {
            "sv": "Skapa en översikt över namn, roller, organisationer och hur de hänger ihop.",
            "en": "Create an overview of names, roles, organizations, and how they connect.",
        },
        "writing_format": "article",
        "writing_tone": "formal",
        "writing_length": "short",
    },
    "write_article_draft": {
        "title": {"sv": "Skriv artikelutkast", "en": "Write article draft"},
        "summary_focus": {
            "sv": "Lyft fram de viktigaste nyhetsvärdena, kärnpåståendena, vad som är dokumenterat och vilka delar som behöver ytterligare verifiering.",
            "en": "Highlight the strongest news value, core claims, what is documented, and which parts need more verification.",
        },
        "summary_style": "investigative",
        "summary_target_pages": 5,
        "summary_words_per_page": 280,
        "writing_brief": {
            "sv": "Skriv ett artikelutkast med tydlig ingress, bakgrund, centrala belägg, relevanta citat och en avslutande lista över sådant som måste verifieras innan publicering.",
            "en": "Write an article draft with a clear lede, background, core evidence, relevant quotes, and a closing list of what must be verified before publication.",
        },
        "writing_format": "article",
        "writing_tone": "investigative",
        "writing_length": "long",
    },
}

LANGUAGES = {
    "sv": {
        "title": "Insikt – Journalist-AI",
        "welcome": "Välkommen till Insikt",
        "howto": "Ladda upp filer i sidofältet och chatta sedan med dem nedan. Assistenten kan svara på frågor, sammanfatta och analysera dina dokument. Alla svar är baserade på dina dokument när möjligt, med källhänvisning.",
        "upload": "Ladda upp dokument",
        "process_btn": "Bearbeta dokument",
        "process_kb_btn": "Bygg kunskapsbas",
        "processing": "Bearbetar... Vänligen vänta. Inga andra åtgärder är möjliga just nu.",
        "success_docs": "{} dokument laddade ({} stycken).",
        "error_no_docs": "Ladda upp dokument först.",
        "knowledge_ready": "Kunskapsbas redo.",
        "chunks_loaded": "{} stycken inlästa från {} dokument.",
        "chat_title": "Chatta med dina dokument",
        "chat_input": "Ställ en fråga...",
        "sources": "Källor",
        "summarize_title": "Sammanfattningsverktyg",
        "summarize_btn": "Generera sammanfattning",
        "focus": "Fokus / instruktioner",
        "target_pages": "Målsidor för sammanfattning",
        "density": "Ord per sida",
        "style": "Utmatningsstil",
        "refine_btn": "Använd förfinad metod (bättre för långa dokument)",
        "summarizing": "Sammanfattar...",
        "success_summary": "Sammanfattning klar.",
        "analysis_title": "Analysverktyg",
        "ner_extract": "Extrahera enheter",
        "timeline": "Skapa tidslinje",
        "keywords": "Extrahera nyckelord",
        "sentiment": "Analysera sentiment",
        "export_title": "Exportera",
        "bias_check": "Kontrollera partiskhet",
        "translate": "Översätt till",
        "session_title": "Session",
        "save_session": "Spara session",
        "load_session": "Ladda session",
        "settings": "Inställningar",
        "language": "Språk",
        "device": "Processor (GPU/CPU)",
        "device_auto": "Auto – använd GPU om tillgängligt, annars CPU",
        "device_cuda": "GPU – tvinga GPU (kräver NVIDIA GPU med CUDA)",
        "device_cpu": "CPU – tvinga CPU (långsammare, men fungerar alltid)",
        "device_current": "Aktuell processor: {}",
        "device_warning_cuda_unavailable": "Varning: Ingen GPU hittades. Kör på CPU.",
        "device_info": "GPU är mycket snabbare för AI-arbete, men kräver ett kompatibelt NVIDIA-kort med CUDA installerat. CPU fungerar alltid men är långsammare.",
        "ollama_gpu_note": "Ollama måste konfigureras separat för GPU. Se Ollamas dokumentation.",
        "summary_in_progress": "En sammanfattning pågår.",
        "summary_will_continue": "Sammanfattningen fortsätter på originalspråket.",
        "error_ollama": "Kunde inte ansluta till Ollama. Kontrollera att Ollama körs (```ollama serve```) och att modellen '{}' är nedladdad.",
        "error_gpu_memory": "GPU-minnet är otillräckligt. Försök med färre eller kortare dokument, eller kör på CPU.",
        "error_pdf_corrupt": "En eller flera PDF-filer är skadade eller oläsbara.",
        "embedding_model": "Inbäddningsmodell",
        "embedding_model_info": "BGE-modeller ger bättre semantisk förståelse än MiniLM",
        "hero_subtitle": "Lokalt AI-stöd för granskande journalistik, manus och redaktionellt arbete.",
        "hero_badge_local": "100% lokalt",
        "hero_badge_private": "Privat",
        "hero_badge_bilingual": "Svenska/English",
        "upload_help": "Välj PDF, DOCX, TXT eller MD",
        "llm_model": "Språkmodell",
        "llm_model_info": "Välj balans mellan hastighet och kvalitet.",
        "chunking": "Segmentering",
        "chunking_info": "Smart segmentering ger bättre sammanhang.",
        "summary_settings": "Sammanfattningsinställningar",
        "summary_estimate": "Uppskattad längd",
        "summary_help": "Bygg en sammanfattning med tydlig struktur och källor.",
        "chat_help": "Ställ frågor om dokumenten eller be om förslag.",
        "writing_title": "Skrivstudio",
        "writing_brief": "Uppdrag / brief",
        "writing_placeholder": "Beskriv vad du vill skapa och för vem...",
        "writing_role": "Roll",
        "writing_format": "Format",
        "writing_tone": "Ton",
        "writing_length": "Längd",
        "writing_use_sources": "Använd dokument som källor",
        "writing_generate": "Skapa manus",
        "writing_result": "Utkast",
        "writing_help": "Använd dokumenten som faktabas när det går.",
        "writing_sources": "Källor använda",
        "writing_pipeline": "Dokumentär-pipeline",
        "writing_pipeline_help": "Skapa en disposition, scenlista och slutligt manus i tre steg.",
        "progress_reading": "Läser dokument",
        "progress_chunking": "Delar upp i segment",
        "progress_indexing": "Bygger kunskapsbas",
        "progress_embedding": "Skapar semantiska vektorer",
        "progress_complete": "Klart.",
        "stage_reading": "Läser dokument {} av {}",
        "stage_chunking": "Delar upp dokument i {} segment",
        "stage_embedding": "Skapar vektorer för {} segment",
        "setup_section": "Kom igång",
        "documents_section": "Dokument",
        "saved_work_section": "Sparat arbete",
        "models_section": "Modeller och prestanda",
        "system_status_section": "Systemstatus",
        "selected_files": "{} filer valda",
        "ready_to_build": "Nästa steg: bygg kunskapsbasen så att appen kan chatta, sammanfatta och analysera filerna.",
        "no_docs_hint": "1. Ladda upp filer. 2. Klicka på Bygg kunskapsbas.",
        "uploaded_sources": "Uppladdade källor",
    },
    "en": {
        "title": "Insikt – Journalist AI",
        "welcome": "Welcome to Insikt",
        "howto": "Upload files in the sidebar, then chat with them below. The assistant can answer questions, summarize, and analyze your documents. All responses are grounded in your documents when possible, with sources cited.",
        "upload": "Upload Documents",
        "process_btn": "Process Documents",
        "process_kb_btn": "Build Knowledge Base",
        "processing": "Processing... Please wait. No other actions are possible at this time.",
        "success_docs": "{} docs loaded ({} chunks).",
        "error_no_docs": "Please upload documents first.",
        "knowledge_ready": "Knowledge base ready.",
        "chunks_loaded": "{} chunks loaded from {} documents.",
        "chat_title": "Chat with Your Documents",
        "chat_input": "Ask a question...",
        "sources": "Sources",
        "summarize_title": "Summarization Tool",
        "summarize_btn": "Generate Summary",
        "focus": "Focus / instructions",
        "target_pages": "Target pages for summary",
        "density": "Words per page",
        "style": "Output style",
        "refine_btn": "Use refine method (better for long docs)",
        "summarizing": "Summarizing...",
        "success_summary": "Summary complete.",
        "analysis_title": "Analysis Tools",
        "ner_extract": "Extract Entities",
        "timeline": "Generate Timeline",
        "keywords": "Extract Keywords",
        "sentiment": "Analyze Sentiment",
        "export_title": "Export",
        "bias_check": "Bias Check",
        "translate": "Translate to",
        "session_title": "Session",
        "save_session": "Save Session",
        "load_session": "Load Session",
        "settings": "Settings",
        "language": "Language",
        "device": "Processing device",
        "device_auto": "Auto – use GPU if available, otherwise CPU",
        "device_cuda": "GPU – force GPU (requires NVIDIA GPU with CUDA)",
        "device_cpu": "CPU – force CPU (slower, but always works)",
        "device_current": "Current device: {}",
        "device_warning_cuda_unavailable": "Warning: No GPU found. Running on CPU.",
        "device_info": "GPU is much faster for AI workloads, but requires a compatible NVIDIA card with CUDA installed. CPU works everywhere but is slower.",
        "ollama_gpu_note": "Ollama must be configured separately for GPU. See Ollama documentation.",
        "summary_in_progress": "A summary is in progress.",
        "summary_will_continue": "The summary will continue in the original language.",
        "error_ollama": "Could not connect to Ollama. Please ensure Ollama is running (```ollama serve```) and the model '{}' is downloaded.",
        "error_gpu_memory": "GPU memory insufficient. Try with fewer or shorter documents, or run on CPU.",
        "error_pdf_corrupt": "One or more PDF files are corrupted or unreadable.",
        "embedding_model": "Embedding Model",
        "embedding_model_info": "BGE models provide better semantic understanding than MiniLM",
        "hero_subtitle": "Local AI support for investigative journalism, scripts, and editorial work.",
        "hero_badge_local": "100% local",
        "hero_badge_private": "Private",
        "hero_badge_bilingual": "Swedish/English",
        "upload_help": "Select PDF, DOCX, TXT, or MD",
        "llm_model": "Language Model",
        "llm_model_info": "Choose the balance between speed and quality.",
        "chunking": "Chunking",
        "chunking_info": "Smart chunking preserves context in long documents.",
        "summary_settings": "Summary settings",
        "summary_estimate": "Estimated length",
        "summary_help": "Build a structured summary with clear sourcing.",
        "chat_help": "Ask questions about the documents or request ideas.",
        "writing_title": "Writing Studio",
        "writing_brief": "Assignment / brief",
        "writing_placeholder": "Describe what you want to create and for whom...",
        "writing_role": "Role",
        "writing_format": "Format",
        "writing_tone": "Tone",
        "writing_length": "Length",
        "writing_use_sources": "Use documents as sources",
        "writing_generate": "Create script",
        "writing_result": "Draft",
        "writing_help": "Use documents as a factual base whenever possible.",
        "writing_sources": "Sources used",
        "writing_pipeline": "Documentary pipeline",
        "writing_pipeline_help": "Generate outline, scene list, and final script in three steps.",
        "progress_reading": "Reading documents",
        "progress_chunking": "Chunking documents",
        "progress_indexing": "Building knowledge base",
        "progress_embedding": "Creating semantic vectors",
        "progress_complete": "Complete.",
        "stage_reading": "Reading document {} of {}",
        "stage_chunking": "Chunking into {} segments",
        "stage_embedding": "Creating vectors for {} segments",
        "setup_section": "Get started",
        "documents_section": "Documents",
        "saved_work_section": "Saved work",
        "models_section": "Models and performance",
        "system_status_section": "System status",
        "selected_files": "{} files selected",
        "ready_to_build": "Next step: build the knowledge base so the app can chat with, summarize, and analyze the files.",
        "no_docs_hint": "1. Upload files. 2. Click Build Knowledge Base.",
        "uploaded_sources": "Uploaded sources",
    }
}


# -------------------------------------------------------------------
# Helper functions (now deviceâ€‘aware)
# -------------------------------------------------------------------
def get_text(key):
    lang = st.session_state.get("lang", "sv")
    return cleaned_ui_text(LANGUAGES[lang].get(key, key))

def resolve_device(choice):
    """Return actual device string based on user choice."""
    if choice == "auto":
        return "cuda" if torch.cuda.is_available() else "cpu"
    elif choice == "cuda":
        if torch.cuda.is_available():
            return "cuda"
        else:
            # Fallback with warning
            st.warning(get_text("device_warning_cuda_unavailable"))
            return "cpu"
    else:  # cpu
        return "cpu"

# We'll cache resources with a dependency on device choice so they reload when device changes.
@st.cache_resource(show_spinner=False)
def load_llm(_device_choice, _model_key=None):  # device_choice is not used directly but forces cache invalidation
    model_key = _model_key or st.session_state.get("llm_model", DEFAULT_LLM_MODEL)
    installed_models = get_installed_ollama_models()
    resolved_model = resolve_installed_ollama_model(model_key, installed_models)
    if installed_models and resolved_model and resolved_model != model_key:
        model_key = resolved_model
        st.session_state.llm_model = resolved_model
    elif installed_models and not resolved_model:
        fallback = "llama3.2:latest" if "llama3.2:latest" in installed_models else installed_models[0]
        st.warning(
            (
                f"Vald modell '{model_key}' finns inte installerad. Byter till '{fallback}'."
                if st.session_state.get("lang", "sv") == "sv"
                else f"Selected model '{model_key}' is not installed. Falling back to '{fallback}'."
            )
        )
        st.session_state.llm_model = fallback
        model_key = fallback
    try:
        return ChatOllama(model=model_key, temperature=0.3, num_predict=2048)
    except Exception as e:
        st.error(get_text("error_ollama").format(model_key))
        st.stop()

@st.cache_resource(show_spinner=False)
def load_embeddings(_device_choice, _embedding_model_key=None):
    device = resolve_device(_device_choice)
    # Get the embedding model key from session state or use default
    model_key = _embedding_model_key or st.session_state.get("embedding_model", DEFAULT_EMBEDDING_MODEL)
    model_info = EMBEDDING_MODELS.get(model_key, EMBEDDING_MODELS[DEFAULT_EMBEDDING_MODEL])
    # Extract the actual model name from the dictionary
    model_name = model_info["model_name"] if isinstance(model_info, dict) else model_info
    try:
        return HuggingFaceEmbeddings(
            model_name=model_name,
            model_kwargs={'device': device}
        )
    except RuntimeError as e:
        if "out of memory" in str(e).lower():
            st.error(get_text("error_gpu_memory"))
            st.stop()
        raise

@st.cache_resource(show_spinner=False)
def load_ner_pipeline(_device_choice):
    device = resolve_device(_device_choice)
    device_id = 0 if device == "cuda" else -1
    try:
        return pipeline("ner", model="dslim/bert-base-NER", aggregation_strategy="simple", device=device_id)
    except Exception as e:
        st.warning("NER pipeline could not be loaded. Feature disabled." if st.session_state.lang=="en" else "NER pipeline kunde inte laddas.")
        return None

@st.cache_resource(show_spinner=False)
def load_sentiment_pipeline(_device_choice):
    device = resolve_device(_device_choice)
    device_id = 0 if device == "cuda" else -1
    try:
        return pipeline("sentiment-analysis", model="distilbert-base-uncased-finetuned-sst-2-english", device=device_id)
    except Exception as e:
        st.warning("Sentiment pipeline could not be loaded." if st.session_state.lang=="en" else "Sentiment pipeline kunde inte laddas.")
        return None

# -------------------------------------------------------------------
# Rest of the helper functions (unchanged except using device-aware loaders)
# -------------------------------------------------------------------
def load_single_pdf(uploaded_file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name
    try:
        loader = PyPDFLoader(tmp_path)
        pages = loader.load()
        for page in pages:
            page.metadata["source"] = uploaded_file.name
        return pages
    except Exception:
        st.error(get_text("error_pdf_corrupt") + f" ({uploaded_file.name})")
        return []
    finally:
        os.unlink(tmp_path)

def load_single_text_file(uploaded_file):
    try:
        raw = uploaded_file.read()
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("latin-1")
    except Exception as e:
        st.error(f"Text file could not be read: {e}")
        return []
    return [Document(page_content=text, metadata={"source": uploaded_file.name, "page": "1"})]

def load_single_docx(uploaded_file):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(uploaded_file.read())
            tmp_path = tmp.name
        doc = DocxDocument(tmp_path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return [Document(page_content=text, metadata={"source": uploaded_file.name, "page": "1"})]
    except Exception as e:
        st.error(f"DOCX file could not be read: {e}")
        return []
    finally:
        if "tmp_path" in locals() and os.path.exists(tmp_path):
            os.unlink(tmp_path)

def load_single_file(uploaded_file):
    suffix = Path(uploaded_file.name).suffix.lower()
    if suffix == ".pdf":
        return load_single_pdf(uploaded_file)
    if suffix in [".txt", ".md"]:
        return load_single_text_file(uploaded_file)
    if suffix == ".docx":
        return load_single_docx(uploaded_file)
    st.warning(f"Unsupported file type: {suffix}")
    return []

def semantic_chunking(pages, embeddings, status_text, threshold=0.5):
    """
    Semantic chunking using embeddings - splits documents based on 
    semantic similarity rather than fixed character counts.
    """
    from langchain_core.documents import Document
    
    status_text.text("Skapar semantiska segment..." if st.session_state.get("lang","sv")=="sv" else "Creating semantic segments...")
    
    # First, split into sentences using basic text splitting
    sentence_splitter = RecursiveCharacterTextSplitter(
        chunk_size=300,  # Small chunks for sentence-level analysis
        chunk_overlap=50,
        separators=[". ", "! ", "? ", "\n"]
    )
    
    # Get all sentences from all pages
    all_sentences = []
    for page in pages:
        sentences = sentence_splitter.split_text(page.page_content)
        for sent in sentences:
            if sent.strip():
                all_sentences.append({
                    "text": sent,
                    "source": page.metadata.get("source", "Unknown"),
                    "page": page.metadata.get("page", "?")
                })
    
    if not all_sentences:
        # Fallback to standard chunking
        splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
        return splitter.split_documents(pages)
    
    # Generate embeddings for all sentences
    status_text.text("Beräknar semantisk likhet..." if st.session_state.get("lang","sv")=="sv" else "Calculating semantic similarity...")
    
    # Batch embed for efficiency
    texts_to_embed = [s["text"] for s in all_sentences]
    embeddings_matrix = embeddings.embed_documents(texts_to_embed)
    
    # Group sentences into chunks based on similarity
    chunks = []
    current_chunk_texts = []
    current_chunk_sources = set()
    current_chunk_pages = set()
    
    for i, sent_info in enumerate(all_sentences):
        current_chunk_texts.append(sent_info["text"])
        current_chunk_sources.add(sent_info["source"])
        current_chunk_pages.add(str(sent_info["page"]))
        
        # Check if we should start a new chunk
        if len(current_chunk_texts) > 1:
            # Compare with previous sentence
            prev_embed = embeddings_matrix[i-1]
            curr_embed = embeddings_matrix[i]
            
            # Calculate cosine similarity
            import numpy as np
            prev_norm = np.linalg.norm(prev_embed)
            curr_norm = np.linalg.norm(curr_embed)
            if prev_norm > 0 and curr_norm > 0:
                similarity = np.dot(prev_embed, curr_embed) / (prev_norm * curr_norm)
            else:
                similarity = 0
            
            # If similarity is below threshold, finalize current chunk
            if similarity < threshold:
                # Finalize current chunk
                chunk_text = " ".join(current_chunk_texts)
                chunks.append(Document(
                    page_content=chunk_text,
                    metadata={
                        "source": ", ".join(current_chunk_sources),
                        "page": ", ".join(current_chunk_pages)
                    }
                ))
                current_chunk_texts = []
                current_chunk_sources = set()
                current_chunk_pages = set()
    
    # Don't forget the last chunk
    if current_chunk_texts:
        chunk_text = " ".join(current_chunk_texts)
        chunks.append(Document(
            page_content=chunk_text,
            metadata={
                "source": ", ".join(current_chunk_sources),
                "page": ", ".join(current_chunk_pages)
            }
        ))
    
    # If semantic chunking produced too few or too many chunks, adjust
    # Merge small chunks or split very large ones
    final_chunks = []
    min_chunk_size = 200
    
    for chunk in chunks:
        if len(chunk.page_content) < min_chunk_size and final_chunks:
            # Merge small chunk with previous
            final_chunks[-1] = Document(
                page_content=final_chunks[-1].page_content + " " + chunk.page_content,
                metadata=final_chunks[-1].metadata
            )
        elif len(chunk.page_content) > CHUNK_SIZE * 2:
            # Split large chunks
            splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
            split_chunks = splitter.split_text(chunk.page_content)
            for sc in split_chunks:
                final_chunks.append(Document(page_content=sc, metadata=chunk.metadata))
        else:
            final_chunks.append(chunk)
    
    return final_chunks

def process_uploaded_files(uploaded_files, progress_bar, status_text, chunking_strategy=None):
    total_files = len(uploaded_files)
    status_text.text(f"{get_text('progress_reading')} (0/{total_files})")
    
    # Get chunking strategy from session state or parameter
    strategy = chunking_strategy or st.session_state.get("chunking_strategy", DEFAULT_CHUNKING)
    
    pages_list = []
    with ThreadPoolExecutor(max_workers=4) as executor:
        futures = {executor.submit(load_single_file, f): f for f in uploaded_files}
        for i, future in enumerate(as_completed(futures)):
            try:
                pages = future.result()
                pages_list.extend(pages)
            except Exception as e:
                st.error(f"Error loading file: {e}")
            progress_bar.progress((i+1)/total_files)
            status_text.text(f"{get_text('progress_reading')} ({i+1}/{total_files})")
    
    if not pages_list:
        st.error("Inga sidor kunde laddas från filerna." if st.session_state.lang=="sv" else "No pages could be loaded from the files.")
        return []
    
    status_text.text(get_text("progress_chunking"))
    
    # Use the selected chunking strategy
    if strategy == "semantic":
        try:
            # Load embeddings for semantic chunking
            embeddings = load_embeddings(st.session_state.device_choice)
            chunks = semantic_chunking(pages_list, embeddings, status_text)
        except Exception as e:
            # Fallback to fixed chunking if semantic fails
            st.warning(f"Semantisk segmentering misslyckades, anvÃ¤nder standard. Fel: {e}" if st.session_state.get("lang","sv")=="sv" else f"Semantic chunking failed, using standard. Error: {e}")
            splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
            chunks = splitter.split_documents(pages_list)
    else:
        # Standard fixed-size chunking
        splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
        chunks = splitter.split_documents(pages_list)
    
    return chunks

def build_vectorstore(chunks, embeddings, progress_bar, status_text):
    status_text.text(get_text("progress_indexing"))
    texts = [doc.page_content for doc in chunks]
    metadatas = [doc.metadata for doc in chunks]
    try:
        vectorstore = FAISS.from_texts(texts, embeddings, metadatas=metadatas)
    except Exception as e:
        st.error(f"Failed to build vector store: {e}")
        return None
    progress_bar.progress(1.0)
    return vectorstore

# ---------- OPTIMIZED SUMMARIZATION WITH BATCH PROCESSING, MAPREDUCE & CACHING ----------
import hashlib

def generate_doc_hash(docs):
    """Generate a unique hash for the document set to use as cache key."""
    hasher = hashlib.md5()
    for doc in docs:
        text = doc.page_content or ""
        hasher.update(text.encode("utf-8", errors="ignore"))
    return hasher.hexdigest()

def get_cached_summary(doc_hash, focus, style, target_words, use_refine):
    """Retrieve cached summary if available."""
    cache_key = f"{doc_hash}_{focus}_{style}_{target_words}_{use_refine}"
    return st.session_state.get("summary_cache", {}).get(cache_key)

def set_cached_summary(doc_hash, focus, style, target_words, use_refine, summary):
    """Store summary in cache."""
    cache_key = f"{doc_hash}_{focus}_{style}_{target_words}_{use_refine}"
    if "summary_cache" not in st.session_state:
        st.session_state.summary_cache = {}
    st.session_state.summary_cache[cache_key] = summary

class SummaryThread(threading.Thread):
    def __init__(self, docs, llm_model, target_pages, focus, style, words_per_page, lang, use_refine, progress_callback):
        super().__init__()
        self.docs = docs
        self.llm_model = llm_model
        self.target_pages = target_pages
        self.focus = focus
        self.style = style
        self.words_per_page = words_per_page
        self.lang = lang
        self.use_refine = use_refine
        self.progress_callback = progress_callback
        self.result = None
        self.error = None
        self._stop_event = threading.Event()

    def stop(self):
        self._stop_event.set()

    def _check_cancelled(self):
        if self._stop_event.is_set():
            self.error = "cancelled"
            return True
        return False

    def _build_batches(self, max_chars=None, max_chunks=None):
        doc_count = len(self.docs)
        if max_chars is None or max_chunks is None:
            if doc_count >= 400:
                max_chars = 32000
                max_chunks = 20
            elif doc_count >= 120:
                max_chars = 24000
                max_chunks = 16
            else:
                max_chars = 16000
                max_chunks = 10
        batches = []
        current = []
        current_chars = 0
        for doc in self.docs:
            source = doc.metadata.get("source", "Unknown")
            page = doc.metadata.get("page", "?")
            if self.lang == "sv":
                entry = f"[Från {source}, sida {page}]: {doc.page_content}\n\n"
            else:
                entry = f"[From {source}, page {page}]: {doc.page_content}\n\n"
            if current and (current_chars + len(entry) > max_chars or len(current) >= max_chunks):
                batches.append("".join(current))
                current = []
                current_chars = 0
            current.append(entry)
            current_chars += len(entry)
        if current:
            batches.append("".join(current))
        return batches

    def _reduce_summary_group(self, llm, reduce_template, summaries, target_words):
        prompt = reduce_template.format(
            existing_summaries="\n\n".join(summaries),
            target_words=target_words,
            focus=self.focus,
            style=self.style,
        )
        return llm.invoke(prompt).content

    def run(self):
        try:
            target_words = max(150, self.target_pages * self.words_per_page)
            num_predict = max(int(target_words * 1.3), 512)
            llm = ChatOllama(model=self.llm_model, temperature=0.2, num_predict=num_predict)

            # ===== STEP 1: BATCH CHUNKS TOGETHER =====
            batches = self._build_batches()

            if self.lang == "sv":
                map_template = """Du är en undersökande journalist. Skapa en detaljerad sammanfattning av följande textavsnitt med fokus på: {focus}.
Stil: {style}. Inkludera viktiga fakta, namn, datum och sidhänvisningar.
Textavsnitt: {text}

Detaljerad sammanfattning av detta avsnitt:"""

                reduce_template = """Du är en journalist och redaktör. Kombinera följande separata sammanfattningar till en sammanhängande helhet på ungefär {target_words} ord.
Fokus: {focus}. Stil: {style}.
Bevara viktiga fakta, namn, datum och källhänvisningar.

Sammanfattningar att kombinera:
{existing_summaries}

Kombinerad sammanfattning:"""

                final_prompt = f"""Du är en journalist och redaktör. Förfina följande sammanfattning till en slutlig, polerad version på ungefär {target_words} ord.
Stil: {{style}}. Fokus: {{focus}}.
Inkludera källhänvisningar [Källa: filnamn, sida X] där möjligt.

Sammanfattning att förfina:
{{current_summary}}

Slutlig sammanfattning:"""
            else:
                map_template = """You are an investigative journalist. Create a detailed summary of the following text passage, focusing on: {focus}.
Style: {style}. Include key facts, names, dates, and page references.
Text passage: {text}

Detailed summary of this passage:"""

                reduce_template = """You are a journalist and editor. Combine the following separate summaries into a coherent whole of approximately {target_words} words.
Focus: {focus}. Style: {style}.
Preserve important facts, names, dates, and source citations.

Summaries to combine:
{existing_summaries}

Combined summary:"""

                final_prompt = """You are a journalist and editor. Refine the following summary into a final, polished version of approximately {target_words} words.
Style: {style}. Focus: {focus}.
Include citations [Source: filename, page X] where possible.

Summary to refine:
{current_summary}

Final summary:"""

            # ===== STEP 2: MAP PHASE - Process batches =====
            batch_summaries = []
            total_batches = len(batches)

            for batch_idx, batch in enumerate(batches):
                if self._check_cancelled():
                    return
                percentage = int((batch_idx + 1) / total_batches * 70)
                if self.progress_callback:
                    self.progress_callback(
                        "processing",
                        batch_idx + 1,
                        total_batches,
                        percentage,
                        f"Sammanfattar avsnitt {batch_idx + 1} av {total_batches}..." if self.lang == "sv" else f"Summarizing section {batch_idx + 1} of {total_batches}..."
                    )
                try:
                    prompt = map_template.format(focus=self.focus, text=batch, style=self.style)
                    summary = llm.invoke(prompt).content
                    batch_summaries.append(summary)
                except Exception as e:
                    self.error = str(e)
                    return

            # ===== STEP 3: REDUCE PHASE =====
            if self._check_cancelled():
                return
            if self.progress_callback:
                self.progress_callback(
                    "combining",
                    total_batches,
                    total_batches,
                    85,
                    "Kombinerar sammanfattningar..." if self.lang == "sv" else "Combining summaries..."
                )

            if len(batch_summaries) == 1:
                combined_summary = batch_summaries[0]
            else:
                summaries_text = "\\n\\n".join(batch_summaries)
                if len(summaries_text) <= 14000:
                    try:
                        prompt = reduce_template.format(
                            existing_summaries=summaries_text,
                            target_words=target_words,
                            focus=self.focus,
                            style=self.style
                        )
                        combined_summary = llm.invoke(prompt).content
                    except Exception as e:
                        self.error = str(e)
                        return
                else:
                    reduction_round = 1
                    combined_level = batch_summaries
                    while len(combined_level) > 1:
                        if self._check_cancelled():
                            return
                        next_level = []
                        group_size = 4 if len(combined_level) > 6 else 3
                        total_groups = (len(combined_level) + group_size - 1) // group_size
                        for group_idx in range(total_groups):
                            if self._check_cancelled():
                                return
                            start = group_idx * group_size
                            group = combined_level[start:start + group_size]
                            if len(group) == 1:
                                next_level.append(group[0])
                                continue
                            if self.progress_callback:
                                percentage = min(85 + reduction_round * 3, 94)
                                self.progress_callback(
                                    "combining",
                                    group_idx + 1,
                                    total_groups,
                                    percentage,
                                    (
                                        f"Kombinerar delsammanfattningar, omgång {reduction_round} ({group_idx + 1}/{total_groups})..."
                                        if self.lang == "sv"
                                        else f"Combining partial summaries, round {reduction_round} ({group_idx + 1}/{total_groups})..."
                                    ),
                                )
                            try:
                                next_level.append(self._reduce_summary_group(llm, reduce_template, group, target_words))
                            except Exception as e:
                                self.error = str(e)
                                return
                        combined_level = next_level
                        reduction_round += 1
                    combined_summary = combined_level[0]

            # ===== STEP 4: Final polish =====
            if self._check_cancelled():
                return
            if self.use_refine:
                if self.progress_callback:
                    self.progress_callback(
                        "polishing",
                        total_batches,
                        total_batches,
                        95,
                        "Färdigställer..." if self.lang == "sv" else "Finalizing..."
                    )
                try:
                    final_prompt_text = final_prompt.format(
                        current_summary=combined_summary,
                        target_words=target_words,
                        style=self.style,
                        focus=self.focus
                    )
                    final_summary = llm.invoke(final_prompt_text).content
                except Exception as e:
                    self.error = str(e)
                    return
                self.result = final_summary
            else:
                self.result = combined_summary

            # Stricter citation verification for summaries
            if self.docs and not _has_citations(self.result, self.lang):
                sources_list = _build_sources_list(self.docs, self.lang)
                self.result = _citation_fix(self.result, self.lang, llm, sources_list)
                if not _has_citations(self.result, self.lang):
                    warning = "\n\nOBS: Källor kunde inte verifieras." if self.lang == "sv" else "\n\nNote: Citations could not be verified."
                    self.result += warning
        except Exception as e:
            self.error = str(e)


def start_summary(docs, target_pages, focus, style, words_per_page, lang, use_refine):
    st.session_state.last_target_words = target_pages * words_per_page
    st.session_state.last_focus = focus
    st.session_state.last_style = style
    st.session_state.last_use_refine = use_refine

    doc_hash = generate_doc_hash(docs)
    target_words = target_pages * words_per_page

    cached = get_cached_summary(doc_hash, focus, style, target_words, use_refine)
    if cached:
        st.session_state.last_summary = cached
        st.session_state.summary_result = cached
        st.success("Cache anvand." if lang == "sv" else "Cache used.")
        return

    model_key = st.session_state.get("llm_model", DEFAULT_LLM_MODEL)
    thread = SummaryThread(docs, model_key, target_pages, focus, style, words_per_page, lang, use_refine, update_summary_progress)
    st.session_state.summary_thread = thread
    st.session_state.summary_progress = 0
    st.session_state.summary_total = len(docs)
    st.session_state.summary_running = True
    st.session_state.summary_result = None
    st.session_state.summary_error = None
    thread.start()


def update_summary_progress(stage, current, total, percentage, message):
    """Enhanced progress callback with detailed stage information."""
    # Initialize session state variables if they don't exist (for background thread safety)
    if "summary_stage" not in st.session_state:
        st.session_state.summary_stage = stage
    else:
        st.session_state.summary_stage = stage
    
    if "summary_current_batch" not in st.session_state:
        st.session_state.summary_current_batch = current
    else:
        st.session_state.summary_current_batch = current
        
    if "summary_total_batches" not in st.session_state:
        st.session_state.summary_total_batches = total
    else:
        st.session_state.summary_total_batches = total
        
    if "summary_percentage" not in st.session_state:
        st.session_state.summary_percentage = percentage
    else:
        st.session_state.summary_percentage = percentage
    
    # Add to log if message changed
    if message:
        # Initialize log if it doesn't exist
        if "summary_stages_log" not in st.session_state:
            st.session_state.summary_stages_log = []
            
        log_entry = {
            "stage": stage,
            "message": message,
            "timestamp": time.time(),
            "percentage": percentage
        }
        if not st.session_state.summary_stages_log or st.session_state.summary_stages_log[-1]["message"] != message:
            st.session_state.summary_stages_log.append(log_entry)
    
        # Keep log manageable - max 50 entries
        if len(st.session_state.summary_stages_log) > 50:
            st.session_state.summary_stages_log = st.session_state.summary_stages_log[-50:]

def check_summary_status():
    if st.session_state.get("summary_running", False):
        thread = st.session_state.get("summary_thread")
        if thread and not thread.is_alive():
            if thread.error:
                st.session_state.summary_error = thread.error
                if thread.error == "cancelled":
                    st.session_state.summary_stage = "cancelled"
                else:
                    st.session_state.summary_stage = "error"
            else:
                st.session_state.summary_result = thread.result
                st.session_state.last_summary = thread.result
                st.session_state.summary_stage = "complete"
                st.session_state.summary_percentage = 100
                
                # Store in cache for future use
                if st.session_state.docs and thread.result:
                    doc_hash = generate_doc_hash(st.session_state.docs)
                    target_words = st.session_state.get("last_target_words", 1500)
                    focus = st.session_state.get("last_focus", "")
                    style = st.session_state.get("last_style", "neutral")
                    use_refine = st.session_state.get("last_use_refine", True)
                    set_cached_summary(doc_hash, focus, style, target_words, use_refine, thread.result)
                    
            st.session_state.summary_running = False
            st.session_state.summary_thread = None
            st.rerun()

# ---------- Chat with RAG ----------
def retrieve_context(query, vectorstore, k=7):
    try:
        return vectorstore.similarity_search(query, k=k)
    except Exception as e:
        st.error(f"SÃ¶kning misslyckades: {e}" if st.session_state.lang=="sv" else f"Search failed: {e}")
        return []

def create_chat_prompt(history, context_docs, query, lang):
    if lang == "sv":
        system_template = """
Du är en hjälpsam och ärlig AI-assistent. Du kan använda användarens uppladdade dokument för att svara på frågor.
När du använder information från dokumenten MÅSTE du ange källa (filnamn och sidnummer) i svaret.
Om frågan inte rör dokumenten eller om du inte hittar relevant information kan du svara utifrån allmän kunskap, men säg tydligt att du inte använder dokumenten.
Var alltid sanningsenlig och hitta inte på information.

Kontext från dokument:
{context}

Tidigare konversation:
{history}

Användare: {query}
Assistent:"""
    else:
        system_template = """
You are a helpful, honest AI assistant. You can access the user's uploaded documents to answer questions.
When you use information from the documents, you MUST cite the source (filename and page number) in your response.
If the user's question is not related to the documents or you don't find relevant information, you can answer based on your general knowledge, but clearly indicate that you are not using the documents.
Always be truthful and do not make up information.

Context from documents:
{context}

Conversation history:
{history}

User: {query}
Assistant:"""

    context_str = ""
    if context_docs:
        for doc in context_docs:
            src = doc.metadata.get("source", "Unknown")
            page = doc.metadata.get("page", "?")
            context_str += f"[KÃ¤lla: {src}, sida {page}]: {doc.page_content}\n\n" if lang == "sv" else f"[Source: {src}, page {page}]: {doc.page_content}\n\n"
    else:
        context_str = "Inga dokument tillgÃ¤ngliga eller ingen relevant information hittades." if lang == "sv" else "No documents available or no relevant information found."

    history_str = ""
    for msg in history[-10:]:
        role = "AnvÃ¤ndare" if isinstance(msg, HumanMessage) else "Assistent"
        if lang == "en":
            role = "User" if isinstance(msg, HumanMessage) else "Assistant"
        history_str += f"{role}: {msg.content}\n"

    prompt = system_template.format(context=context_str, history=history_str, query=query)
    return prompt

def chat_with_docs(query, history, vectorstore, llm, lang):
    # Use adaptive retrieval - analyze query complexity to determine k
    k = analyze_query_complexity(query)
    context_docs = retrieve_context(query, vectorstore, k=k) if vectorstore else []
    prompt = create_chat_prompt(history, context_docs, query, lang)
    try:
        response = llm.invoke(prompt).content
    except Exception as e:
        response = f"Fel vid generering av svar: {e}" if lang == "sv" else f"Error generating answer: {e}"
        return response, context_docs
    
    # Verify citations exist in retrieved context
    verified_response, citation_issues = verify_citations(response, context_docs, lang)
    
    # If there are citation issues, we can either regenerate or flag the issue
    # For now, we return the verified response with any warnings
    return verified_response, context_docs


# ---------- Writing Studio ----------
def create_writing_prompt(brief, role_key, format_key, tone_key, length_key, lang, context_docs):
    role_label = WRITING_ROLES.get(role_key, {}).get(lang, role_key)
    format_label = WRITING_FORMATS.get(format_key, {}).get(lang, format_key)
    tone_label = WRITING_TONES.get(tone_key, {}).get(lang, tone_key)
    length_words = WRITING_LENGTHS.get(length_key, {}).get("words", 800)

    context_str = ""
    if context_docs:
        for doc in context_docs[:6]:
            src = doc.metadata.get("source", "Unknown")
            page = doc.metadata.get("page", "?")
            if lang == "sv":
                context_str += f"[Källa: {src}, sida {page}]: {doc.page_content}\n\n"
            else:
                context_str += f"[Source: {src}, page {page}]: {doc.page_content}\n\n"

    if lang == "sv":
        prompt = f"""Du är en {role_label}. Skapa ett utkast i formatet {format_label}.
Ton: {tone_label}. Mållängd: cirka {length_words} ord.

Uppdrag:
{brief}

När du använder fakta från dokument, inkludera källhänvisningar [Källa: filnamn, sida X].
Om fakta saknas, markera antaganden tydligt.

Källor:
{context_str if context_str else 'Inga källutdrag tillgängliga.'}

Leverera:
1) Kort disposition
2) Utkast/manus med tydlig struktur
3) Lista över saknade källor om något behöver verifieras"""
    else:
        prompt = f"""You are a {role_label}. Create a draft in the format {format_label}.
Tone: {tone_label}. Target length: about {length_words} words.

Assignment:
{brief}

When using facts from documents, include citations [Source: filename, page X].
If you lack facts, clearly mark assumptions.

Sources:
{context_str if context_str else 'No source snippets available.'}

Deliver:
1) Short outline
2) Draft/script with clear structure
3) Source needs list if anything is missing"""

    return prompt

def _has_citations(text, lang):
    if lang == "sv":
        return bool(re.search(r"\[Källa:\s*[^\]]+\]", text))
    return bool(re.search(r"\[Source:\s*[^\]]+\]", text))


def _citation_fix(text, lang, llm, sources_list):
    if not sources_list:
        return text
    if lang == "sv":
        prompt = f"""Du är en redaktör. Lägg till tydliga källhänvisningar i texten.
Använd endast denna lista över källor:
{sources_list}

Text:
{text}

Uppdaterad text med källor:"""
    else:
        prompt = f"""You are an editor. Add clear citations to the text.
Use only this list of sources:
{sources_list}

Text:
{text}

Updated text with citations:"""
    try:
        return llm.invoke(prompt).content
    except Exception:
        return text


def _build_sources_list(docs, lang, limit=15):
    if not docs:
        return ""
    items = []
    seen = set()
    for doc in docs:
        src = doc.metadata.get("source", "Unknown")
        page = doc.metadata.get("page", "?")
        key = f"{src}:{page}"
        if key in seen:
            continue
        seen.add(key)
        label = f"{src}, sida {page}" if lang == "sv" else f"{src}, page {page}"
        items.append(label)
        if len(items) >= limit:
            break
    return "\n".join(items)

def generate_writing(brief, role_key, format_key, tone_key, length_key, lang, vectorstore, llm, use_sources=True, use_pipeline=False):
    context_docs = retrieve_context(brief, vectorstore, k=6) if vectorstore and use_sources else []
    prompt = create_writing_prompt(brief, role_key, format_key, tone_key, length_key, lang, context_docs)
    sources_list = _build_sources_list(context_docs, lang)
    try:
        if use_pipeline and format_key == "documentary":
            outline_prompt = prompt + ("\n\nSteg 1: Skapa en tydlig disposition." if lang == "sv" else "\n\nStep 1: Create a clear outline.")
            outline = llm.invoke(outline_prompt).content
            scene_prompt = ("\n\nSteg 2: Skapa en scenlista baserat på dispositionen:\n" if lang == "sv" else "\n\nStep 2: Create a scene list based on the outline:\n") + outline
            scenes = llm.invoke(scene_prompt).content
            script_prompt = ("\n\nSteg 3: Skriv slutligt manus baserat på scenlistan:\n" if lang == "sv" else "\n\nStep 3: Write the final script based on the scene list:\n") + scenes
            response = llm.invoke(script_prompt).content
            response = f"{outline}\n\n{scenes}\n\n{response}"
        else:
            response = llm.invoke(prompt).content
    except Exception as e:
        response = f"Fel vid generering: {e}" if lang == "sv" else f"Error generating text: {e}"

    if use_sources and context_docs and not _has_citations(response, lang):
        response = _citation_fix(response, lang, llm, sources_list)
        if not _has_citations(response, lang):
            warning = "\n\nOBS: Källor kunde inte verifieras." if lang == "sv" else "\n\nNote: Citations could not be verified."
            response += warning
    return response, context_docs


# ---------- Adaptive Retrieval: Query Complexity Analysis ----------
def analyze_query_complexity(query: str) -> int:
    """
    Analyze query complexity and return optimal k value for retrieval.
    
    Simple queries: k=3-4
    Medium complexity: k=5-7
    Complex analytical: k=8-10
    """
    query_lower = query.lower()
    
    # Complexity indicators
    complexity_score = 0
    
    # Length-based complexity (longer queries often need more context)
    word_count = len(query.split())
    if word_count > 15:
        complexity_score += 2
    elif word_count > 8:
        complexity_score += 1
    
    # Analytical keywords indicate complex queries
    analytical_keywords = [
        'compare', 'analyze', 'explain', 'why', 'how', 'difference',
        'relationship', 'impact', 'effect', 'cause', 'result', 'consequence',
        'jÃ¤mfÃ¶r', 'analysera', 'fÃ¶rklara', 'varfÃ¶r', 'hur', 'skillnad',
        'relation', 'pÃ¥verkan', 'effekt', 'orsak', 'resultat', 'konsekvens'
    ]
    for keyword in analytical_keywords:
        if keyword in query_lower:
            complexity_score += 2
    
    # Multiple entities/concepts (question words)
    question_indicators = ['who', 'what', 'when', 'where', 'which', 'whom', 'whose',
                           'vem', 'vad', 'nÃ¤r', 'var', 'vilken', 'vilket', 'vilka']
    for indicator in question_indicators:
        if indicator in query_lower:
            complexity_score += 1
    
    # Comparison queries need more context
    if any(word in query_lower for word in ['vs', 'versus', 'compared', 'eller', 'jÃ¤mfÃ¶rt']):
        complexity_score += 2
    
    # Map complexity score to k value
    if complexity_score <= 2:
        return 3  # Simple query
    elif complexity_score <= 5:
        return 5  # Medium complexity
    else:
        return 8  # Complex analytical query


# ---------- Citation Verification ----------
def extract_citations_from_response(response: str) -> list:
    """
    Extract citations from the LLM response.
    Looks for patterns like [Source: filename, page X] or [KÃ¤lla: filename, sida X]
    """
    citations = []
    
    # Pattern for English: [Source: filename, page X]
    en_pattern = r'\[Source:\s*([^,\]]+),\s*page\s*(\d+)\]'
    # Pattern for Swedish: [KÃ¤lla: filename, sida X]
    sv_pattern = r'\[KÃ¤lla:\s*([^,\]]+),\s*sida\s*(\d+)\]'
    
    # Also match bare citations like (source, page 5) or [source, p.5]
    bare_pattern = r'\[?([A-Za-z0-9_\-\.]+)[,\s]+(?:page|p\.?|sida|s\.)\s*(\d+)\]?'
    
    import re
    for match in re.finditer(en_pattern, response, re.IGNORECASE):
        citations.append({
            'source': match.group(1).strip(),
            'page': match.group(2).strip()
        })
    
    for match in re.finditer(sv_pattern, response, re.IGNORECASE):
        citations.append({
            'source': match.group(1).strip(),
            'page': match.group(2).strip()
        })
    
    for match in re.finditer(bare_pattern, response, re.IGNORECASE):
        citations.append({
            'source': match.group(1).strip(),
            'page': match.group(2).strip()
        })
    
    return citations


def verify_citations(response: str, context_docs: list, lang: str) -> tuple:
    """
    Verify that citations in the response actually exist in the retrieved context.
    Returns (verified_response, issues)
    """
    if not context_docs:
        return response, ["No context available to verify citations"]
    
    # Extract citations from response
    citations = extract_citations_from_response(response)
    
    if not citations:
        # No citations found - this is fine, just return as-is
        return response, []
    
    # Build a set of available sources from context
    available_sources = {}
    for doc in context_docs:
        source = doc.metadata.get("source", "Unknown")
        page = doc.metadata.get("page", "?")
        if source not in available_sources:
            available_sources[source] = set()
        available_sources[source].add(str(page))
    
    # Verify each citation
    issues = []
    verified_citations = []
    
    for citation in citations:
        source = citation['source']
        page = citation['page']
        
        # Check if source exists in context
        source_found = False
        for avail_source in available_sources:
            # Partial match (e.g., "report.pdf" matches "report_final.pdf")
            if source.lower() in avail_source.lower() or avail_source.lower() in source.lower():
                source_found = True
                # Check if page exists (optional - some sources might not have page numbers)
                if page in available_sources[avail_source] or page == "?":
                    verified_citations.append(citation)
                else:
                    # Page not found but source exists - might still be valid
                    verified_citations.append(citation)
                break
        
        if not source_found:
            issues.append(f"Citation refers to '{source}' which was not found in retrieved context")
    
    # If there are issues, add a warning to the response
    if issues:
        warning_msg = ""
        if lang == "sv":
            warning_msg = "\n\nObs: Vissa källhänvisningar kunde inte verifieras i dokumentkontexten."
        else:
            warning_msg = "\n\nNote: Some citations could not be verified in the document context."
        
        response = response + warning_msg
    
    return response, issues

# ---------- Analysis tools ----------
def extract_entities(text, ner_pipeline):
    if not ner_pipeline:
        return {}
    results = ner_pipeline(text[:10000])
    entities = {"PERSON": set(), "ORG": set(), "LOC": set()}
    for ent in results:
        if ent['entity_group'] in entities:
            entities[ent['entity_group']].add(ent['word'])
    return entities

def extract_timeline(docs):
    date_pattern = r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}\b'
    timeline = []
    for doc in docs:
        text = doc.page_content
        for match in re.finditer(date_pattern, text, re.IGNORECASE):
            date = match.group()
            start = max(0, match.start() - 100)
            end = min(len(text), match.end() + 100)
            context = text[start:end].replace('\n', ' ')
            timeline.append({
                "date": date,
                "context": context,
                "source": doc.metadata.get("source", "Unknown"),
                "page": doc.metadata.get("page", "?")
            })
    timeline.sort(key=lambda x: x["date"])
    return timeline

def bias_check(summary, llm, lang):
    if lang == "sv":
        prompt = f"""
Du Ã¤r en kritisk redaktÃ¶r. Granska fÃ¶ljande sammanfattning fÃ¶r eventuell partiskhet, saknad motbevisning eller Ã¶verdriven tilltro till en enda kÃ¤lla.
Lista eventuella problem du hittar och fÃ¶reslÃ¥ vad en balanserad sammanfattning skulle innehÃ¥lla.

Sammanfattning:
{summary}

Kritik:
"""
    else:
        prompt = f"""
You are a critical editor. Review the following summary for potential bias, missing counterâ€‘evidence, or overâ€‘reliance on a single source.
List any issues you find, and suggest what a balanced summary would include.

Summary:
{summary}

Critique:
"""
    try:
        return llm.invoke(prompt).content
    except Exception as e:
        return f"Bias check failed: {e}"

def translate_text(text, target_lang, llm, source_lang):
    prompt = f"Ã–versÃ¤tt fÃ¶ljande text till {target_lang}:\n\n{text}" if source_lang == "sv" else f"Translate the following text to {target_lang}:\n\n{text}"
    try:
        return llm.invoke(prompt).content
    except Exception as e:
        return f"Translation failed: {e}"

def extract_keywords(text, top_n=10, language="en"):
    try:
        kw_extractor = yake.KeywordExtractor(lan=language, top=top_n)
        keywords = kw_extractor.extract_keywords(text)
        return [kw for kw, score in keywords]
    except Exception as e:
        return [f"Error: {e}"]

def analyze_sentiment(text, pipe):
    if not pipe:
        return {"label": "N/A", "score": 0.0}
    try:
        return pipe(text[:512])[0]
    except Exception as e:
        return {"label": "Error", "score": 0.0}

# ---------- Export ----------
def export_text(text): return text.encode()
def export_docx(text):
    doc = DocxDocument()
    doc.add_paragraph(text)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        with open(tmp.name, "rb") as f: data = f.read()
    os.unlink(tmp.name); return data
def export_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in text.split('\n'):
        pdf.multi_cell(0, 10, line)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        pdf.output(tmp.name)
        with open(tmp.name, "rb") as f: data = f.read()
    os.unlink(tmp.name); return data
def export_markdown(text): return text.encode()

# -------------------------------------------------------------------
# UI Custom CSS (unchanged)
# -------------------------------------------------------------------
def set_custom_css(basic_mode: bool = False):
    base_font = "1.04rem" if basic_mode else "0.98rem"
    button_padding = "0.8rem 1.1rem" if basic_mode else "0.55rem 1.1rem"
    button_font = "1rem" if basic_mode else "0.95rem"
    input_font = "1rem" if basic_mode else "0.95rem"
    css = """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Sans:wght@400;500;600&family=Source+Serif+4:wght@500;600&display=swap');
    :root {
        --ink: #0f172a;
        --muted: #475569;
        --paper: #fbf7f2;
        --paper-2: #f4efe8;
        --accent: #d97706;
        --accent-2: #0f766e;
        --card: #ffffff;
        --border: #e5e7eb;
        --shadow: 0 12px 30px rgba(15, 23, 42, 0.08);
    }
    html, body, [class*="css"] {{
        font-family: 'IBM Plex Sans', sans-serif;
        font-size: __BASE_FONT__;
    }}
    h1, h2, h3, .hero-title { font-family: 'Source Serif 4', serif; }
    .stApp {
        background:
            radial-gradient(1200px 600px at 10% -10%, #fff7ed 0%, transparent 60%),
            radial-gradient(900px 500px at 90% -20%, #ecfeff 0%, transparent 55%),
            var(--paper);
        color: var(--ink);
    }
    [data-testid="stSidebar"] {
        background: var(--paper-2);
        border-right: 1px solid var(--border);
    }
    .hero {
        background: rgba(255,255,255,0.7);
        border: 1px solid var(--border);
        border-radius: 18px;
        padding: 1.5rem 1.8rem;
        box-shadow: var(--shadow);
    }
    .hero-title { font-size: 2.2rem; font-weight: 600; margin-bottom: 0.4rem; }
    .hero-subtitle { color: var(--muted); font-size: 1.05rem; }
    .chip {
        display: inline-block;
        background: #fff;
        border: 1px solid var(--border);
        border-radius: 999px;
        padding: 0.25rem 0.7rem;
        margin-right: 0.4rem;
        font-size: 0.85rem;
        color: var(--muted);
    }
    .card {
        background: var(--card);
        border-radius: 18px;
        padding: 1.2rem 1.4rem;
        margin: 0.8rem 0;
        box-shadow: var(--shadow);
        border: 1px solid var(--border);
    }
    .processing-overlay {
        position: fixed;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        background: rgba(255,255,255,0.9);
        backdrop-filter: blur(3px);
        z-index: 999;
        display: flex;
        justify-content: center;
        align-items: center;
        flex-direction: column;
    }
    .processing-box {
        background: white;
        border-radius: 24px;
        padding: 2rem;
        box-shadow: 0 20px 25px -5px rgba(0,0,0,0.1), 0 10px 10px -5px rgba(0,0,0,0.04);
        border: 1px solid var(--border);
        max-width: 400px;
        text-align: center;
    }
    .chat-message {
        padding: 1rem;
        border-radius: 12px;
        margin-bottom: 0.5rem;
        max-width: 80%;
    }
    .user-message {
        background: var(--accent-2);
        color: #ffffff;
        margin-left: auto;
    }
    .assistant-message {
        background: #fffaf5;
        color: var(--ink);
        margin-right: auto;
        border: 1px solid var(--border);
    }
    .source-box {
        background: #fff;
        border-left: 3px solid var(--accent);
        padding: 0.5rem 1rem;
        margin: 0.5rem 0;
        font-size: 0.9rem;
        border-radius: 0 10px 10px 0;
    }
    .stButton>button {{
        background: var(--accent);
        color: #ffffff;
        border: none;
        border-radius: 10px;
        font-weight: 600;
        padding: __BUTTON_PADDING__;
        font-size: __BUTTON_FONT__;
        min-height: 2.8rem;
        transition: all 0.2s ease;
    }}
    .stButton>button:hover:not(:disabled) {{
        background: #b45309;
        box-shadow: 0 10px 18px rgba(15, 23, 42, 0.12);
        transform: translateY(-1px);
    }}
    .stButton>button:disabled {{
        background: #94a3b8;
        cursor: not-allowed;
        opacity: 0.6;
    }}
    .stTextInput input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"], .stMultiSelect div[data-baseweb="select"] {{
        font-size: __INPUT_FONT__;
    }}
    .stCheckbox label, .stRadio label {{
        font-size: __INPUT_FONT__;
    }}
    .stProgress > div > div > div > div { background-color: var(--accent); }
    .info-box {
        background: #fff;
        border: 1px solid var(--border);
        border-radius: 12px;
        padding: 1rem;
        margin: 0.5rem 0;
    }
    .success-box {
        background: #ecfdf3;
        border: 1px solid #86efac;
        border-radius: 12px;
        padding: 1rem;
        margin: 0.5rem 0;
    }
    .warning-box {
        background: #fff7ed;
        border: 1px solid #fdba74;
        border-radius: 12px;
        padding: 1rem;
        margin: 0.5rem 0;
    }
    h1, h2, h3 { font-weight: 600; }
    .stExpander { border: 1px solid var(--border); border-radius: 14px; margin-bottom: 1rem; }
    .chat-container {
        display: flex;
        flex-direction: column;
        height: 60vh;
        overflow-y: auto;
        padding: 1rem;
        border: 1px solid var(--border);
        border-radius: 16px;
        margin-bottom: 1rem;
        background: rgba(255,255,255,0.7);
    }
    .stChatInput {
        margin-top: 1rem;
    }
    div[data-baseweb="tab"] {
        font-weight: 600;
        color: var(--muted);
        padding: 0.6rem 1rem;
    }
    div[data-baseweb="tab"][aria-selected="true"] {
        color: var(--ink);
        border-bottom: 3px solid var(--accent);
    }
    [data-testid="stPopover"] > button {
        width: 2rem;
        height: 2rem;
        min-height: 2rem;
        border-radius: 999px;
        padding: 0;
        background: #ffffff;
        color: var(--accent-2);
        border: 1px solid var(--border);
        font-weight: 700;
        box-shadow: none;
    }
    [data-testid="stPopover"] > button:hover {
        background: #f8fafc;
        transform: none;
        box-shadow: 0 6px 16px rgba(15, 23, 42, 0.08);
    }
    div[data-baseweb="popover"] {
        max-width: 420px !important;
    }
    .help-block-title {
        font-size: 0.85rem;
        text-transform: uppercase;
        letter-spacing: 0.04em;
        color: var(--accent-2);
        margin-bottom: 0.4rem;
    }
    .chat-shell {
        background: rgba(255,255,255,0.62);
        border: 1px solid var(--border);
        border-radius: 18px;
        padding: 1rem;
        box-shadow: var(--shadow);
    }
    [data-testid="stChatMessage"] {
        max-width: 860px;
        margin-left: auto;
        margin-right: auto;
        padding-top: 0.35rem;
        padding-bottom: 0.35rem;
    }
    [data-testid="stChatMessageContent"] {
        border-radius: 18px;
        padding: 1rem 1.15rem;
        line-height: 1.65;
    }
    [data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-user"]) {
        justify-content: flex-end;
    }
    [data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-user"]) [data-testid="stChatMessageContent"] {
        background: linear-gradient(135deg, #0f766e 0%, #115e59 100%);
        border: 1px solid #0f766e;
        color: #ffffff;
        margin-left: auto;
        box-shadow: 0 12px 30px rgba(15, 118, 110, 0.18);
    }
    [data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-assistant"]) [data-testid="stChatMessageContent"] {
        background: #fffaf5;
        border: 1px solid var(--border);
        margin-right: auto;
        box-shadow: 0 12px 28px rgba(15, 23, 42, 0.06);
    }
    [data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-user"]) [data-testid="stMarkdownContainer"] p,
    [data-testid="stChatMessage"]:has([data-testid="chatAvatarIcon-user"]) [data-testid="stMarkdownContainer"] li {
        color: #ffffff;
    }
    .chat-notes {
        margin-top: 0.8rem;
        padding: 0.75rem 0.9rem;
        border-radius: 12px;
        background: #fff7ed;
        border: 1px solid #fdba74;
    }
    .chat-sources {
        margin-top: 0.8rem;
        border-top: 1px solid var(--border);
        padding-top: 0.8rem;
    }
    .citation-register {
        margin-top: 0.8rem;
        padding: 0.85rem 1rem;
        background: #f8fafc;
        border: 1px solid var(--border);
        border-radius: 12px;
    }
    .citation-register strong {
        display: block;
        margin-bottom: 0.45rem;
    }
    .session-thumb {
        background: linear-gradient(180deg, #fffdf8 0%, #fff7ed 100%);
        border: 1px solid var(--border);
        border-radius: 14px;
        padding: 0.9rem 1rem;
        margin: 0.7rem 0 0.45rem 0;
        box-shadow: var(--shadow);
    }
    .session-thumb-title {
        font-weight: 700;
        margin-bottom: 0.35rem;
        color: var(--ink);
    }
    .session-thumb-meta {
        color: var(--muted);
        font-size: 0.88rem;
        margin-bottom: 0.45rem;
    }
    .session-thumb-preview {
        color: #334155;
        font-size: 0.93rem;
        line-height: 1.5;
    }
    </style>
    """
    css = (
        css.replace("__BASE_FONT__", base_font)
        .replace("__BUTTON_PADDING__", button_padding)
        .replace("__BUTTON_FONT__", button_font)
        .replace("__INPUT_FONT__", input_font)
    )
    st.markdown(css, unsafe_allow_html=True)


SummaryThread = ModularSummaryThread
generate_doc_hash = modular_generate_doc_hash


@st.cache_data(show_spinner=False)
def get_startup_checks(model_key: str, embedding_key: str):
    return run_startup_checks(model_key, embedding_key)


def set_status(status_placeholder, message):
    status_placeholder.text(cleaned_ui_text(message))


def refresh_source_options():
    raw_pages = st.session_state.get("raw_pages") or []
    st.session_state.available_sources = sorted({doc.metadata.get("source", "Unknown") for doc in raw_pages})


def parse_tags(raw_value: str) -> list[str]:
    return [part.strip() for part in raw_value.split(",") if part.strip()]


def build_issue_context(build_meta: dict) -> dict:
    return {
        "app_build": build_meta.get("label", ""),
        "language": st.session_state.get("lang", "sv"),
        "device_choice": st.session_state.get("device_choice", "auto"),
        "llm_model": st.session_state.get("llm_model", DEFAULT_LLM_MODEL),
        "embedding_model": st.session_state.get("embedding_model", DEFAULT_EMBEDDING_MODEL),
        "loaded_sources": st.session_state.get("available_sources", []),
        "doc_count": len(st.session_state.get("raw_pages") or []),
        "chat_messages": len(st.session_state.get("chat_history") or []),
        "summary_available": bool(st.session_state.get("last_summary")),
        "writing_available": bool(st.session_state.get("writing_result")),
    }


def open_folder_in_file_manager(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)
    if os.name == "nt":
        os.startfile(str(path))  # type: ignore[attr-defined]
        return
    if sys.platform == "darwin":
        subprocess.Popen(["open", str(path)])
        return
    subprocess.Popen(["xdg-open", str(path)])


def save_tester_issue(build_meta: dict) -> dict:
    payload = {
        "reporter_name": st.session_state.get("issue_reporter_name", "").strip(),
        "severity": st.session_state.get("issue_severity", "Medium"),
        "area": st.session_state.get("issue_area", "General"),
        "title": st.session_state.get("issue_title", "").strip(),
        "what_happened": st.session_state.get("issue_what_happened", "").strip(),
        "expected": st.session_state.get("issue_expected", "").strip(),
        "steps": st.session_state.get("issue_steps", "").strip(),
        "work_context": st.session_state.get("issue_work_context", "").strip(),
        "app_version_label": build_meta.get("label", ""),
        "app_context": build_issue_context(build_meta),
    }
    return save_issue_report(FEEDBACK_ROOT, payload)


def format_slot_timestamp(timestamp: str) -> str:
    if not timestamp:
        return ""
    cleaned = timestamp.replace("Z", "+00:00")
    try:
        return datetime.fromisoformat(cleaned).strftime("%Y-%m-%d %H:%M")
    except ValueError:
        return timestamp[:16]


def parse_multiline_list(raw_value: str) -> list[str]:
    return [line.strip() for line in (raw_value or "").splitlines() if line.strip()]


def serialize_case_board() -> dict:
    return {
        "notes": st.session_state.get("case_board_notes", ""),
        "people": parse_multiline_list(st.session_state.get("case_board_people", "")),
        "dates": parse_multiline_list(st.session_state.get("case_board_dates", "")),
        "angles": parse_multiline_list(st.session_state.get("case_board_angles", "")),
        "excerpts": st.session_state.get("case_board_excerpts", []),
    }


def apply_case_board(payload: dict | None):
    payload = payload or {}
    st.session_state.case_board_notes = payload.get("notes", "")
    st.session_state.case_board_people = "\n".join(payload.get("people", []))
    st.session_state.case_board_dates = "\n".join(payload.get("dates", []))
    st.session_state.case_board_angles = "\n".join(payload.get("angles", []))
    st.session_state.case_board_excerpts = payload.get("excerpts", [])


def add_case_board_excerpt(source: str, page: str, excerpt: str):
    if not excerpt:
        return
    item = {"source": str(source or ""), "page": str(page or ""), "excerpt": excerpt.strip()}
    current = st.session_state.get("case_board_excerpts", [])
    if any(entry.get("source") == item["source"] and entry.get("page") == item["page"] and entry.get("excerpt") == item["excerpt"] for entry in current):
        return
    st.session_state.case_board_excerpts = [item] + current[:24]


def get_reporter_template_options(lang: str) -> list[str]:
    return [REPORTER_TEMPLATES[key]["title"][lang] for key in REPORTER_TEMPLATES]


def template_key_from_label(label: str, lang: str) -> str | None:
    for key, template in REPORTER_TEMPLATES.items():
        if template["title"][lang] == label:
            return key
    return None


def apply_reporter_template(template_key: str, lang: str):
    template = REPORTER_TEMPLATES.get(template_key)
    if not template:
        return
    st.session_state.summary_focus = template["summary_focus"][lang]
    st.session_state.summary_style = template["summary_style"]
    st.session_state.summary_target_pages = template["summary_target_pages"]
    st.session_state.summary_words_per_page = template["summary_words_per_page"]
    st.session_state.writing_brief = template["writing_brief"][lang]
    st.session_state.writing_format = template["writing_format"]
    st.session_state.writing_tone = template["writing_tone"]
    st.session_state.writing_length = template["writing_length"]
    st.session_state.reporter_template = template_key


def reset_document_dependent_state(clear_chat: bool = True):
    if clear_chat:
        st.session_state.chat_history = []
    st.session_state.quote_candidates = []
    st.session_state.comparison_results = []
    st.session_state.comparison_query = ""
    st.session_state.comparison_last_run = ""
    st.session_state.writing_sources = []
    st.session_state.selected_preview_source = ""
    st.session_state.selected_preview_page = ""
    st.session_state.preview_excerpt = ""
    st.session_state.preview_origin = "manual"
    st.session_state.recent_citations = []
    st.session_state.preview_bookmarks = []


def ingest_uploaded_files(uploaded_files):
    st.session_state.processing = True
    try:
        embeddings = load_embeddings(st.session_state.device_choice, st.session_state.get("embedding_model", DEFAULT_EMBEDDING_MODEL))
        progress = st.progress(0.0)
        status = st.empty()
        fingerprint, raw_pages, chunks, ingest_stats = cached_process_uploaded_files(
            uploaded_files,
            embeddings=embeddings,
            cache_root=CACHE_ROOT,
            chunking_strategy=st.session_state.get("chunking_strategy", DEFAULT_CHUNKING),
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            status_callback=lambda msg: set_status(status, msg),
            progress_callback=lambda value: progress.progress(min(value, 1.0)),
            error_callback=lambda file_name: st.error(f"{get_text('error_pdf_corrupt')} ({file_name})"),
            use_file_cache=st.session_state.get("performance_mode", False),
        )
        vectorstore = build_or_load_vectorstore(fingerprint, chunks, embeddings, CACHE_ROOT, status_callback=lambda msg: set_status(status, msg))
        st.session_state.raw_pages = raw_pages
        st.session_state.docs = chunks
        st.session_state.vectorstore = vectorstore
        st.session_state.doc_fingerprint = fingerprint
        st.session_state.last_ingest_stats = ingest_stats
        reset_document_dependent_state()
        refresh_source_options()
        st.success(get_text("success_docs").format(len(uploaded_files), len(chunks)))
    finally:
        st.session_state.processing = False


def rebuild_from_current_pages():
    raw_pages = st.session_state.get("raw_pages") or []
    if not raw_pages:
        st.session_state.docs = None
        st.session_state.vectorstore = None
        st.session_state.doc_fingerprint = ""
        refresh_source_options()
        return
    embeddings = load_embeddings(st.session_state.device_choice, st.session_state.get("embedding_model", DEFAULT_EMBEDDING_MODEL))
    chunks = rechunk_pages(
        raw_pages,
        embeddings,
        chunking_strategy=st.session_state.get("chunking_strategy", DEFAULT_CHUNKING),
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    fingerprint = modular_generate_doc_hash(raw_pages)
    st.session_state.docs = chunks
    st.session_state.vectorstore = build_or_load_vectorstore(fingerprint, chunks, embeddings, CACHE_ROOT)
    st.session_state.doc_fingerprint = fingerprint
    reset_document_dependent_state()
    refresh_source_options()


def remove_source(source_name: str):
    st.session_state.raw_pages = [doc for doc in (st.session_state.get("raw_pages") or []) if doc.metadata.get("source") != source_name]
    rebuild_from_current_pages()


def apply_loaded_slot(payload: dict):
    st.session_state.docs = payload.get("docs") or []
    st.session_state.raw_pages = payload.get("raw_pages") or []
    st.session_state.chat_history = payload.get("chat_history") or []
    st.session_state.last_summary = payload.get("last_summary", "")
    st.session_state.vectorstore = payload.get("vectorstore")
    st.session_state.lang = payload.get("lang", st.session_state.get("lang", "sv"))
    st.session_state.doc_fingerprint = payload.get("fingerprint", "")
    reset_document_dependent_state(clear_chat=False)
    apply_case_board(payload.get("case_board"))
    refresh_source_options()


def copy_block(text: str, key: str):
    escaped = text.replace("\\", "\\\\").replace("`", "\\`").replace("</", "<\\/")
    components.html(
        f"""
        <button onclick="navigator.clipboard.writeText(`{escaped}`)" style="padding:0.45rem 0.8rem;border:1px solid #d6d3d1;border-radius:8px;background:#fff;cursor:pointer;">
            Copy
        </button>
        """,
        height=42,
    )


def render_page_header(title: str, caption: str, help_title: str, help_body: str):
    title_col, help_col = st.columns([20, 1])
    with title_col:
        st.markdown(f"### {title}")
    with help_col:
        with st.popover("i", help=help_title):
            st.markdown(f"**{help_title}**")
            st.markdown(help_body)
    if caption:
        st.caption(caption)


def split_assistant_content(content: str):
    parts = [part.strip() for part in content.split("\n\n") if part.strip()]
    body_parts = []
    notes = []
    for part in parts:
        if part.startswith("Obs:") or part.startswith("OBS:") or part.startswith("Note:"):
            notes.append(part)
        else:
            body_parts.append(part)
    return "\n\n".join(body_parts).strip(), notes


def build_citation_link(source: str, page: str) -> str:
    return f"?cite_source={quote(str(source))}&cite_page={quote(str(page))}"


def select_preview_target(source: str, page: str = "", excerpt: str = "", origin: str = "manual"):
    st.session_state.selected_preview_source = str(source or "")
    st.session_state.selected_preview_page = str(page or "")
    st.session_state.preview_excerpt = excerpt or ""
    st.session_state.preview_origin = origin
    if origin in {"citation", "register", "snippet"} and source:
        target = {
            "source": str(source),
            "page": str(page or ""),
            "excerpt": excerpt or "",
            "origin": origin,
        }
        recent = st.session_state.get("recent_citations", [])
        recent = [item for item in recent if not (item.get("source") == target["source"] and item.get("page") == target["page"])]
        recent.insert(0, target)
        st.session_state.recent_citations = recent[:8]
        st.session_state.last_citation_target = f"{target['source']}:{target['page']}"


def add_preview_bookmark(source: str, page: str, excerpt: str = ""):
    if not source:
        return
    bookmark = {"source": str(source), "page": str(page or ""), "excerpt": excerpt or ""}
    bookmarks = st.session_state.get("preview_bookmarks", [])
    if any(item.get("source") == bookmark["source"] and item.get("page") == bookmark["page"] for item in bookmarks):
        return
    bookmarks.append(bookmark)
    st.session_state.preview_bookmarks = bookmarks


def prettify_citations_for_display(text: str, lang: str):
    pattern = r"\[(Källa|Source):\s*([^,\]]+),\s*(sida|page)\s*([^\]]+)\]"
    source_order = []
    source_map = {}

    def replacer(match):
        source = match.group(2).strip()
        page = match.group(4).strip()
        key = (source, page)
        if key not in source_map:
            source_map[key] = len(source_order) + 1
            source_order.append(key)
        index = source_map[key]
        return f"[[{index}]]({build_citation_link(source, page)})"

    cleaned = re.sub(pattern, replacer, text, flags=re.IGNORECASE)
    register = []
    for source, page in source_order:
        register.append(
            {
                "index": source_map[(source, page)],
                "source": source,
                "page": page,
                "label": f"[{source_map[(source, page)]}] {source}, {'sida' if lang == 'sv' else 'page'} {page}",
            }
        )
    return cleaned, register


def render_citation_register(register: list[dict], lang: str, key_prefix: str):
    if not register:
        return
    st.markdown('<div class="citation-register">', unsafe_allow_html=True)
    st.markdown("**Källregister**" if lang == "sv" else "**Citation Register**")
    for item in register:
        col_a, col_b = st.columns([1, 12])
        with col_a:
            if st.button(f"[{item['index']}]", key=f"{key_prefix}-cite-{item['index']}"):
                select_preview_target(item["source"], str(item["page"]), origin="register")
                st.rerun()
        with col_b:
            st.markdown(item["label"])
    st.markdown("</div>", unsafe_allow_html=True)


def render_source_snippets(sources, lang: str, key_prefix: str):
    if not sources:
        return
    if sources and isinstance(sources[0], dict):
        sources = records_to_docs(sources)
    with st.expander("Visa källutdrag" if lang == "sv" else "Show source excerpts", expanded=False):
        for source_index, doc in enumerate(sources[:5]):
            src = doc.metadata.get("source", "Unknown")
            page = doc.metadata.get("page", "?")
            label = f"s.{page}" if lang == "sv" else f"p.{page}"
            st.markdown(
                f'<div class="source-box"><b>{safe_html_fragment(str(src))}</b> ({safe_html_fragment(label)})<br>{safe_html_fragment(doc.page_content[:220])}...</div>',
                unsafe_allow_html=True,
            )
            if st.button("Visa källa" if lang == "sv" else "Open source", key=f"{key_prefix}-source-{source_index}"):
                select_preview_target(src, str(page), excerpt=doc.page_content[:500], origin="snippet")
                st.rerun()


def render_confidence_banner(confidence: dict, lang: str):
    level = confidence.get("level", "needs_review")
    if level == "well_supported":
        st.success(f"{confidence.get('label')}: {confidence.get('reason')}")
    elif level == "partly_supported":
        st.info(f"{confidence.get('label')}: {confidence.get('reason')}")
    else:
        st.warning(f"{confidence.get('label')}: {confidence.get('reason')}")
    citation_count = int(confidence.get("citation_count", 0) or 0)
    source_count = int(confidence.get("source_count", 0) or 0)
    if citation_count > 0:
        st.caption(
            (
                f"Källhänvisningar: {citation_count} | Källdokument: {source_count}"
                if lang == "sv"
                else f"Citations: {citation_count} | Source documents: {source_count}"
            )
        )


def get_llm_option_info(model_key: str) -> dict:
    if model_key in LLM_MODELS:
        return LLM_MODELS[model_key]
    base_key = model_key.split(":", 1)[0]
    if base_key in LLM_MODELS:
        return LLM_MODELS[base_key]
    return {
        "display_name": model_key,
        "display_name_en": model_key,
        "description": "Installerad lokal modell.",
        "description_en": "Installed local model.",
        "speed": "",
        "quality": "",
        "memory": "",
    }


def get_available_llm_options(installed_models: list[str]) -> list[str]:
    return installed_models if installed_models else list(LLM_MODELS.keys())


def describe_system_profile(profile: dict, lang: str) -> str:
    cpu = profile.get("cpu_count", 0)
    ram = profile.get("ram_gb", 0.0)
    if profile.get("gpu_available"):
        gpu_name = profile.get("gpu_name", "GPU")
        vram = profile.get("vram_gb", 0.0)
        return (
            f"{cpu} CPU-kärnor | {ram} GB RAM | {gpu_name} ({vram} GB VRAM)"
            if lang == "sv"
            else f"{cpu} CPU cores | {ram} GB RAM | {gpu_name} ({vram} GB VRAM)"
        )
    return (
        f"{cpu} CPU-kärnor | {ram} GB RAM | Ingen CUDA-GPU hittades"
        if lang == "sv"
        else f"{cpu} CPU cores | {ram} GB RAM | No CUDA GPU detected"
    )


def recommended_startup_preset(recommendations: list[dict]) -> dict | None:
    for key in ("balanced", "fast", "best"):
        for preset in recommendations:
            if preset.get("key") == key:
                return preset
    return recommendations[0] if recommendations else None


def auto_fix_setup(primary_preset: dict, installed_models: list[str], lang: str) -> list[str]:
    actions = []
    missing_packages = get_missing_python_packages()
    if missing_packages:
        try:
            installed_package_names = install_missing_python_packages()
            package_list = ", ".join(installed_package_names)
            actions.append(
                f"Installerade saknade Python-paket: {package_list}."
                if lang == "sv"
                else f"Installed missing Python packages: {package_list}."
            )
        except Exception as exc:
            actions.append(
                f"Kunde inte installera alla Python-paket automatiskt: {exc}"
                if lang == "sv"
                else f"Could not install all Python packages automatically: {exc}"
            )
    if st.session_state.get("device_choice") == "cuda" and not torch.cuda.is_available():
        st.session_state.device_choice = "auto"
        actions.append("Återställde processorn till Auto." if lang == "sv" else "Reset processing device to Auto.")
    if st.session_state.get("embedding_model") not in EMBEDDING_MODELS:
        st.session_state.embedding_model = primary_preset["embedding_model"]
        actions.append(
            f"Valde inbäddningsprofilen {primary_preset['embedding_model']}."
            if lang == "sv"
            else f"Selected embedding profile {primary_preset['embedding_model']}."
        )
    target_model = primary_preset["llm_model"]
    resolved_current = resolve_installed_ollama_model(st.session_state.get("llm_model", DEFAULT_LLM_MODEL), installed_models)
    if resolved_current and resolved_current != st.session_state.get("llm_model"):
        st.session_state.llm_model = resolved_current
        actions.append(
            f"Justerade språkmodellen till installerade {resolved_current}."
            if lang == "sv"
            else f"Adjusted the language model to installed model {resolved_current}."
        )
    if installed_models and target_model not in installed_models:
        resolved_target = resolve_installed_ollama_model(target_model, installed_models)
        if resolved_target:
            target_model = resolved_target
    if target_model != st.session_state.get("llm_model"):
        st.session_state.llm_model = target_model
        actions.append(
            f"Valde rekommenderad språkmodell {target_model}."
            if lang == "sv"
            else f"Selected recommended language model {target_model}."
        )
    if st.session_state.get("embedding_model") != primary_preset["embedding_model"]:
        st.session_state.embedding_model = primary_preset["embedding_model"]
        actions.append(
            f"Valde rekommenderad sökprofil {primary_preset['embedding_model']}."
            if lang == "sv"
            else f"Selected recommended search profile {primary_preset['embedding_model']}."
        )
    if not shutil.which("tesseract"):
        actions.append(
            f"OCR behöver också Tesseract-kommandot i PATH. {get_tesseract_install_hint()}"
            if lang == "sv"
            else f"OCR also needs the Tesseract CLI on PATH. {get_tesseract_install_hint()}"
        )
    get_startup_checks.clear()
    load_llm.clear()
    load_embeddings.clear()
    return actions


def render_system_check(check: dict, lang: str):
    friendly_names = {
        "Python dependencies": "Appens komponenter" if lang == "sv" else "App components",
        "GPU readiness": "Grafikprocessor" if lang == "sv" else "GPU",
        "Ollama": "Ollama-tjänst" if lang == "sv" else "Ollama service",
        "OCR readiness": "Skannade PDF-filer" if lang == "sv" else "Scanned PDFs",
        "Selected LLM": "Vald språkmodell" if lang == "sv" else "Selected language model",
        "Embedding profile": "Sökprofil" if lang == "sv" else "Search profile",
    }
    title = friendly_names.get(check["name"], check["name"])
    line = f"{title}: {cleaned_ui_text(check['message'])}"
    if check["status"] == "warning":
        st.warning(line)
    elif check["status"] == "ok":
        st.success(line)
    else:
        st.info(line)


def pull_ollama_model(model_name: str):
    result = subprocess.run(["ollama", "pull", model_name], capture_output=True, text=True, timeout=1800000, check=False)
    get_startup_checks.clear()
    load_llm.clear()
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or result.stdout.strip() or f"Failed to pull {model_name}")

# -------------------------------------------------------------------
# Main app
# -------------------------------------------------------------------
def main():
    st.set_page_config(page_title=APP_NAME, layout="wide")
    if "basic_mode" not in st.session_state:
        st.session_state.basic_mode = True
    set_custom_css(st.session_state.get("basic_mode", False))

    defaults = {
        "docs": None,
        "raw_pages": [],
        "vectorstore": None,
        "chat_history": [],
        "last_summary": "",
        "lang": "sv",
        "device_choice": "auto",
        "basic_mode": True,
        "performance_mode": False,
        "llm_model": DEFAULT_LLM_MODEL,
        "embedding_model": DEFAULT_EMBEDDING_MODEL,
        "chunking_strategy": DEFAULT_CHUNKING,
        "processing": False,
        "summary_running": False,
        "summary_thread": None,
        "summary_progress": 0,
        "summary_total": 0,
        "summary_result": None,
        "summary_error": None,
        "summary_stage": "idle",
        "summary_stages_log": [],
        "summary_current_batch": 0,
        "summary_total_batches": 0,
        "summary_percentage": 0,
        "summary_cancel_requested": False,
        "summary_focus": "",
        "summary_target_pages": 5,
        "summary_words_per_page": 300,
        "summary_style": "neutral",
        "summary_use_refine": True,
        "reporter_template": "",
        "writing_brief": "",
        "writing_role": "author",
        "writing_format": "documentary",
        "writing_tone": "investigative",
        "writing_length": "medium",
        "writing_use_sources": True,
        "writing_pipeline": False,
        "writing_result": "",
        "writing_sources": [],
        "quote_candidates": [],
        "comparison_query": "",
        "comparison_results": [],
        "comparison_last_run": "",
        "available_sources": [],
        "selected_preview_source": "",
        "selected_preview_page": "",
        "preview_excerpt": "",
        "preview_origin": "manual",
        "preview_bookmarks": [],
        "recent_citations": [],
        "last_citation_target": "",
        "save_slot_name": "",
        "save_slot_folder": "",
        "save_slot_tags": "",
        "session_folder_filter": "All",
        "session_tag_filter": "All",
        "case_board_notes": "",
        "case_board_people": "",
        "case_board_dates": "",
        "case_board_angles": "",
        "case_board_excerpts": [],
        "last_ingest_stats": {},
        "doc_fingerprint": "",
        "issue_reporter_name": "",
        "issue_severity": "Medium",
        "issue_area": "General",
        "issue_title": "",
        "issue_what_happened": "",
        "issue_expected": "",
        "issue_steps": "",
        "issue_work_context": "",
    }
    for key, val in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = val

    check_summary_status()
    refresh_source_options()
    lang = st.session_state.lang
    query_params = st.query_params
    cite_source = query_params.get("cite_source")
    cite_page = query_params.get("cite_page")
    if cite_source:
        target = f"{cite_source}:{cite_page or ''}"
        if target != st.session_state.get("last_citation_target", ""):
            select_preview_target(cite_source, str(cite_page or ""), origin="citation")
    system_profile = get_system_profile()
    build_meta = get_build_metadata()
    checks = get_startup_checks(st.session_state.get("llm_model", DEFAULT_LLM_MODEL), st.session_state.get("embedding_model", DEFAULT_EMBEDDING_MODEL))
    readiness_state, readiness_text = app_readiness_label(st.session_state.get("docs"), st.session_state.get("processing", False))
    installed_models = get_installed_ollama_models()
    cache_stats = get_cache_stats(CACHE_ROOT)
    available_llm_options = get_available_llm_options(installed_models)
    recommendations = get_model_recommendations(installed_models, system_profile, list(LLM_MODELS.keys()))
    recent_reports = list_issue_reports(FEEDBACK_ROOT, limit=12)
    primary_preset = recommended_startup_preset(recommendations)
    resolved_selected_model = resolve_installed_ollama_model(st.session_state.get("llm_model", DEFAULT_LLM_MODEL), installed_models)
    if installed_models and not resolved_selected_model:
        preferred_model = "llama3.2:latest" if "llama3.2:latest" in installed_models else installed_models[0]
        st.session_state.llm_model = preferred_model
        resolved_selected_model = preferred_model
    selected_llm_value = resolved_selected_model or st.session_state.get("llm_model", DEFAULT_LLM_MODEL)
    if selected_llm_value not in available_llm_options and available_llm_options:
        selected_llm_value = available_llm_options[0]
        st.session_state.llm_model = selected_llm_value

    with st.sidebar:
        st.markdown(f"## {get_text('title')}")
        st.caption(build_meta.get("label", ""))
        current_device = resolve_device(st.session_state.device_choice)
        if readiness_state == "ready":
            st.success("Kunskapsbas redo" if lang == "sv" else readiness_text)
        elif readiness_state == "processing":
            st.info("Bearbetning pågår. Appen arbetar i bakgrunden och kan ta en stund beroende på filstorlek och modell." if lang == "sv" else "Processing is in progress. The app is still working and may take a while depending on file size and model.")
        else:
            st.warning("Ingen kunskapsbas laddad" if lang == "sv" else readiness_text)
        st.caption(get_text("device_current").format(current_device.upper()))

        with st.expander(get_text("setup_section"), expanded=True):
            st.caption("Snabbstart" if lang == "sv" else "Quick start")
            st.caption(describe_system_profile(system_profile, lang))
            lang_options = ["sv", "en"]
            selected_lang = st.selectbox(
                get_text("language"),
                options=lang_options,
                index=lang_options.index(st.session_state.get("lang", "sv")),
                format_func=lambda x: "Svenska" if x == "sv" else "English",
            )
            if selected_lang != st.session_state.lang:
                st.session_state.lang = selected_lang
                st.rerun()
            basic_mode = st.toggle(
                "Enkelt läge" if lang == "sv" else "Basic mode",
                value=st.session_state.get("basic_mode", True),
                help="Förenklar språket och döljer fler tekniska val." if lang == "sv" else "Uses simpler language and hides more technical options.",
            )
            if basic_mode != st.session_state.get("basic_mode", True):
                st.session_state.basic_mode = basic_mode
                st.rerun()
            if basic_mode:
                st.caption(
                    "Appen väljer rekommenderade inställningar åt dig. Öppna avancerat läge om du vill finjustera."
                    if lang == "sv"
                    else "The app will use recommended settings for you. Open advanced mode if you want manual control."
                )
            else:
                selected_device = st.selectbox(get_text("device"), options=["auto", "cuda", "cpu"], format_func=lambda x: get_text(f"device_{x}"))
                if selected_device != st.session_state.device_choice:
                    st.session_state.device_choice = selected_device
                    st.rerun()
                st.caption(get_text("device_info"))
                st.caption(get_text("ollama_gpu_note"))
            if primary_preset:
                st.info(
                    (
                        f"Rekommenderad startprofil: {primary_preset['label']} med {primary_preset['llm_model']}."
                        if lang == "sv"
                        else f"Recommended starter profile: {primary_preset['label']} with {primary_preset['llm_model']}."
                    )
                )
                if st.button("Fixa vanliga problem automatiskt" if lang == "sv" else "Auto-fix common setup issues", use_container_width=True):
                    messages = auto_fix_setup(primary_preset, installed_models, lang)
                    if messages:
                        for msg in messages:
                            st.success(msg)
                    else:
                        st.success("Inställningarna såg redan bra ut." if lang == "sv" else "The setup already looked good.")
                    st.rerun()
                selected_model_resolved = resolve_installed_ollama_model(st.session_state.get("llm_model", DEFAULT_LLM_MODEL), installed_models)
                if primary_preset["llm_model"] not in installed_models and primary_preset["llm_model"] in LLM_MODELS:
                    install_label = (
                        f"Ladda ner rekommenderad modell ({primary_preset['llm_model']})"
                        if lang == "sv"
                        else f"Download recommended model ({primary_preset['llm_model']})"
                    )
                    if st.button(install_label, use_container_width=True):
                        with st.spinner("Laddar ner modell..." if lang == "sv" else "Downloading model..."):
                            try:
                                pull_ollama_model(primary_preset["llm_model"])
                                st.session_state.llm_model = primary_preset["llm_model"]
                                st.success(
                                    f"Modellen {primary_preset['llm_model']} är installerad."
                                    if lang == "sv"
                                    else f"Model {primary_preset['llm_model']} is installed."
                                )
                                st.rerun()
                            except Exception as exc:
                                st.error(
                                    f"Kunde inte ladda ner modellen: {exc}"
                                    if lang == "sv"
                                    else f"Could not download model: {exc}"
                                )
                elif selected_model_resolved:
                    st.caption(
                        (
                            f"Nuvarande modell kan användas direkt: {selected_model_resolved}."
                            if lang == "sv"
                            else f"Current model is ready to use: {selected_model_resolved}."
                        )
                    )
            with st.expander(get_text("system_status_section"), expanded=False):
                for check in checks:
                    render_system_check(check, lang)
                missing_python_packages = get_missing_python_packages()
                if missing_python_packages:
                    if st.button(
                        "Installera saknade Python-paket" if lang == "sv" else "Install missing Python packages",
                        use_container_width=True,
                    ):
                        with st.spinner("Installerar paket..." if lang == "sv" else "Installing packages..."):
                            try:
                                installed_package_names = install_missing_python_packages()
                                package_list = ", ".join(installed_package_names)
                                st.success(
                                    f"Installerade: {package_list}"
                                    if lang == "sv"
                                    else f"Installed: {package_list}"
                                )
                                st.rerun()
                            except Exception as exc:
                                st.error(
                                    f"Kunde inte installera Python-paketen: {exc}"
                                    if lang == "sv"
                                    else f"Could not install the Python packages: {exc}"
                                )
                if not shutil.which("tesseract"):
                    st.caption(
                        (
                            f"För OCR på skannade PDF-filer behöver du också Tesseract. {get_tesseract_install_hint()}"
                            if lang == "sv"
                            else f"Scanned PDF OCR also needs Tesseract. {get_tesseract_install_hint()}"
                        )
                    )
                if not basic_mode:
                    with st.expander("Tekniska detaljer" if lang == "sv" else "Technical details", expanded=False):
                        for check in checks:
                            st.caption(f"{check['name']}: {cleaned_ui_text(check['message'])}")

        with st.expander(get_text("models_section"), expanded=not basic_mode):
            st.caption("Rekommendationer baserade på den här datorn." if lang == "sv" else "Recommendations based on this computer.")
            for preset in recommendations:
                title = (
                    f"**{preset['label']}**: {preset['llm_model']} + {preset['embedding_model']}"
                )
                st.markdown(title)
                st.caption(
                    (
                        f"{preset['summary']} {'Installerad och klar.' if preset['installed'] else 'Kan väljas manuellt eller laddas ner.'}"
                        if lang == "sv"
                        else f"{preset['summary']} {'Installed and ready.' if preset['installed'] else 'Can be selected manually or downloaded.'}"
                    )
                )
                if st.button(
                    (f"Använd {preset['label']}" if lang == "sv" else f"Use {preset['label']}"),
                    key=f"use-preset-{preset['key']}",
                    use_container_width=True,
                ):
                    st.session_state.llm_model = preset["llm_model"]
                    st.session_state.embedding_model = preset["embedding_model"]
                    get_startup_checks.clear()
                    load_llm.clear()
                    load_embeddings.clear()
                    st.rerun()
            if basic_mode:
                current_model_name = st.session_state.get("llm_model", DEFAULT_LLM_MODEL)
                st.caption(
                    (
                        f"Nuvarande val: {current_model_name} + {st.session_state.get('embedding_model', DEFAULT_EMBEDDING_MODEL)}"
                        if lang == "sv"
                        else f"Current choice: {current_model_name} + {st.session_state.get('embedding_model', DEFAULT_EMBEDDING_MODEL)}"
                    )
                )
            else:
                llm_model = st.selectbox(
                    get_text("llm_model"),
                    options=available_llm_options,
                    index=available_llm_options.index(selected_llm_value) if available_llm_options else 0,
                    format_func=lambda key: concise_model_label(get_llm_option_info(key), lang),
                    help=("Visar installerade Ollama-modeller på den här datorn." if lang == "sv" else "Shows Ollama models installed on this computer."),
                )
                if llm_model != st.session_state.get("llm_model", DEFAULT_LLM_MODEL):
                    st.session_state.llm_model = llm_model
                    st.rerun()
                missing_known_models = [model for model in LLM_MODELS.keys() if model not in installed_models]
                if missing_known_models:
                    download_target = st.selectbox(
                        "Ladda ner modell" if lang == "sv" else "Download model",
                        options=missing_known_models,
                        format_func=lambda key: concise_model_label(get_llm_option_info(key), lang),
                        help="Hämtar modellen via Ollama och gör den tillgänglig i appen." if lang == "sv" else "Downloads the model with Ollama and makes it available in the app.",
                    )
                    if st.button("Ladda ner vald modell" if lang == "sv" else "Download selected model", use_container_width=True):
                        with st.spinner(("Laddar ner modell..." if lang == "sv" else "Downloading model...")):
                            try:
                                pull_ollama_model(download_target)
                                st.success((f"Modellen {download_target} är installerad." if lang == "sv" else f"Model {download_target} is installed."))
                                st.session_state.llm_model = download_target
                                st.rerun()
                            except Exception as exc:
                                st.error((f"Kunde inte ladda ner modellen: {exc}" if lang == "sv" else f"Could not download model: {exc}"))
                embedding_model = st.selectbox(
                    get_text("embedding_model"),
                    options=list(EMBEDDING_MODELS.keys()),
                    index=list(EMBEDDING_MODELS.keys()).index(st.session_state.get("embedding_model", DEFAULT_EMBEDDING_MODEL)),
                    format_func=lambda key: concise_model_label(EMBEDDING_MODELS.get(key, {}), lang),
                )
                if embedding_model != st.session_state.get("embedding_model", DEFAULT_EMBEDDING_MODEL):
                    st.session_state.embedding_model = embedding_model
                    st.rerun()
                chunking = st.selectbox(
                    get_text("chunking"),
                    options=list(CHUNKING_STRATEGIES.keys()),
                    index=list(CHUNKING_STRATEGIES.keys()).index(st.session_state.get("chunking_strategy", DEFAULT_CHUNKING)),
                    format_func=lambda key: CHUNKING_STRATEGIES[key]["display_name_en"] if lang == "en" else CHUNKING_STRATEGIES[key]["display_name"],
                )
                if chunking != st.session_state.get("chunking_strategy", DEFAULT_CHUNKING):
                    st.session_state.chunking_strategy = chunking
                    if st.session_state.get("raw_pages"):
                        rebuild_from_current_pages()
                    st.rerun()

        with st.expander(get_text("documents_section"), expanded=True):
            performance_mode = st.checkbox(
                "Stort dokumentläge" if lang == "sv" else "Large-document mode",
                value=st.session_state.get("performance_mode", False),
                key="performance_mode",
                help="Återanvänder filcache och gör uppdateringar smidigare när bara någon fil ändras."
                if lang == "sv"
                else "Reuses file cache and makes updates smoother when only a few files change.",
            )
            st.caption(
                (
                    f"Cache: {cache_stats['bundle_caches']} buntar, {cache_stats['file_caches']} filcacher"
                    if lang == "sv"
                    else f"Cache: {cache_stats['bundle_caches']} bundles, {cache_stats['file_caches']} file caches"
                )
            )
            ocr_check = next((check for check in checks if check.get("name") == "OCR readiness"), None)
            if ocr_check:
                st.caption(cleaned_ui_text(ocr_check.get("message", "")))
            uploaded_files = st.file_uploader(get_text("upload"), type=["pdf", "docx", "txt", "md"], accept_multiple_files=True, disabled=st.session_state.processing, help=get_text("upload_help"))
            if uploaded_files:
                st.success(get_text("selected_files").format(len(uploaded_files)))
                st.caption(get_text("ready_to_build"))
                preview_names = ", ".join(file.name for file in uploaded_files[:3])
                if len(uploaded_files) > 3:
                    preview_names += f" +{len(uploaded_files) - 3}"
                st.caption(preview_names)
            elif not st.session_state.get("docs"):
                st.info(get_text("no_docs_hint"))
            if uploaded_files and st.button(get_text("process_kb_btn"), disabled=st.session_state.processing, use_container_width=True):
                ingest_uploaded_files(uploaded_files)
                st.rerun()
            ingest_stats = st.session_state.get("last_ingest_stats", {})
            if ingest_stats:
                if ingest_stats.get("bundle_cache_hit"):
                    st.info("Hela dokumentuppsättningen laddades från cache." if lang == "sv" else "The full document set was loaded from cache.")
                elif performance_mode:
                    st.caption(
                        (
                            f"Återanvände {ingest_stats.get('file_cache_hits', 0)} filer från cache, läste om {ingest_stats.get('files_processed', 0)} filer."
                            if lang == "sv"
                            else f"Reused {ingest_stats.get('file_cache_hits', 0)} files from cache and re-read {ingest_stats.get('files_processed', 0)} files."
                        )
                    )
            if st.session_state.get("docs"):
                unique_sources = len(set(doc.metadata.get("source", "Unknown") for doc in st.session_state.get("raw_pages", [])))
                st.info(get_text("chunks_loaded").format(len(st.session_state.docs), unique_sources))
                st.caption(get_text("uploaded_sources"))
            for source in st.session_state.get("available_sources", []):
                col_preview, col_remove = st.columns([3, 1])
                with col_preview:
                    if st.button(source, key=f"preview-source-{source}", use_container_width=True):
                        select_preview_target(source, origin="manual")
                with col_remove:
                    if st.button("X", key=f"remove-source-{source}", use_container_width=True):
                        remove_source(source)
                        st.rerun()
        with st.expander(get_text("saved_work_section"), expanded=False):
            st.text_input("Namn på sparning" if lang == "sv" else "Save name", key="save_slot_name", help="Ge sparningen ett eget namn så att du kan ladda den senare." if lang == "sv" else "Give this save a name so you can load it later.")
            st.text_input("Case-mapp" if lang == "sv" else "Case folder", key="save_slot_folder", help="Till exempel granskning, intervju eller research." if lang == "sv" else "For example investigation, interview, or research.")
            st.text_input("Taggar" if lang == "sv" else "Tags", key="save_slot_tags", help="Separera med kommatecken." if lang == "sv" else "Separate with commas.")
            if st.button("Spara nuvarande läge" if lang == "sv" else "Save current slot", use_container_width=True, disabled=st.session_state.processing, help="Sparar dokument, chatt, sammanfattning och index i en egen sparning." if lang == "sv" else "Saves documents, chat, summary, and index into its own save slot."):
                if st.session_state.get("docs"):
                    slot_name = st.session_state.get("save_slot_name") or f"Save {len(list_save_slots(SAVES_ROOT)) + 1}"
                    slot_id = save_slot(
                        SAVES_ROOT,
                        slot_name,
                        st.session_state.get("docs") or [],
                        st.session_state.get("raw_pages") or [],
                        st.session_state.get("chat_history") or [],
                        st.session_state.get("last_summary", ""),
                        st.session_state.get("lang", "sv"),
                        st.session_state.get("vectorstore"),
                        st.session_state.get("doc_fingerprint", ""),
                        case_folder=st.session_state.get("save_slot_folder", ""),
                        tags=parse_tags(st.session_state.get("save_slot_tags", "")),
                        case_board=serialize_case_board(),
                    )
                    st.success(f"{'Slot sparad' if lang == 'sv' else 'Slot saved'}: {slot_id}")
                else:
                    st.warning("Inget att spara" if lang == "sv" else "Nothing to save")
            save_slots = list_save_slots(SAVES_ROOT)
            if save_slots:
                st.caption("Sparade lägen" if lang == "sv" else "Saved slots")
                folder_options = ["All"] + sorted({slot.get("case_folder", "").strip() or ("No folder" if lang == "sv" else "No folder") for slot in save_slots})
                tag_options = ["All"] + sorted({tag for slot in save_slots for tag in slot.get("tags", [])})
                filter_col_a, filter_col_b = st.columns(2)
                with filter_col_a:
                    selected_folder = st.selectbox("Filtrera mapp" if lang == "sv" else "Folder filter", options=folder_options, key="session_folder_filter")
                with filter_col_b:
                    selected_tag = st.selectbox("Filtrera tagg" if lang == "sv" else "Tag filter", options=tag_options if len(tag_options) > 1 else ["All"], key="session_tag_filter")
                filtered_slots = []
                for slot in save_slots:
                    slot_folder = slot.get("case_folder", "").strip() or ("No folder" if lang == "sv" else "No folder")
                    folder_match = selected_folder == "All" or slot_folder == selected_folder
                    tag_match = selected_tag == "All" or selected_tag in slot.get("tags", [])
                    if folder_match and tag_match:
                        filtered_slots.append(slot)
                grouped_slots = {}
                for slot in filtered_slots:
                    folder_name = slot.get("case_folder", "").strip() or ("Utan mapp" if lang == "sv" else "No folder")
                    grouped_slots.setdefault(folder_name, []).append(slot)
                for folder_name, slots_in_folder in grouped_slots.items():
                    st.markdown(f"**{folder_name}**")
                    for slot in slots_in_folder:
                        tag_text = ", ".join(slot.get("tags", []))
                        meta_parts = [
                            f"{slot.get('doc_count', 0)} {'segment' if lang == 'sv' else 'chunks'}",
                            format_slot_timestamp(slot.get("updated_at", "")),
                        ]
                        if tag_text:
                            meta_parts.append(f"{'Taggar' if lang == 'sv' else 'Tags'}: {tag_text}")
                        preview_source = slot.get("preview_source", "")
                        preview_text = slot.get("preview_text", "")
                        if preview_source:
                            meta_parts.append(preview_source)
                        session_title = safe_html_fragment(str(slot.get('name', slot.get('slot_id'))))
                        session_meta = safe_html_fragment(" | ".join(part for part in meta_parts if part))
                        session_preview = safe_html_fragment(preview_text or ("Ingen förhandsvisning ännu." if lang == "sv" else "No preview yet."))
                        st.markdown(
                            f"""
                            <div class="session-thumb">
                                <div class="session-thumb-title">{session_title}</div>
                                <div class="session-thumb-meta">{session_meta}</div>
                                <div class="session-thumb-preview">{session_preview}</div>
                            </div>
                            """,
                            unsafe_allow_html=True,
                        )
                        col_load, col_delete = st.columns(2)
                        with col_load:
                            if st.button("Ladda" if lang == "sv" else "Load", key=f"load-slot-{slot['slot_id']}", use_container_width=True):
                                embeddings = load_embeddings(st.session_state.device_choice, st.session_state.get("embedding_model", DEFAULT_EMBEDDING_MODEL))
                                apply_loaded_slot(load_slot(SAVES_ROOT, slot["slot_id"], embeddings))
                                st.success("Sparning laddad" if lang == "sv" else "Save loaded")
                                st.rerun()
                        with col_delete:
                            if st.button("Ta bort" if lang == "sv" else "Delete", key=f"delete-slot-{slot['slot_id']}", use_container_width=True):
                                delete_slot(SAVES_ROOT, slot["slot_id"])
                                st.rerun()
            st.caption("All behandling sker lokalt. Sparningar ligger i session_data/saves." if lang == "sv" else "All processing is local. Saves are stored in session_data/saves.")

        with st.expander("Rapportera problem" if lang == "sv" else "Report issue", expanded=False):
            st.caption(
                "Berätta kort vad som gick fel. Appen sparar varje rapport separat så inget skrivs över."
                if lang == "sv"
                else "Describe what went wrong in simple words. The app saves each report separately, so nothing gets overwritten."
            )
            helper_col_a, helper_col_b = st.columns(2)
            with helper_col_a:
                if st.button("Öppna rapportmapp" if lang == "sv" else "Open feedback folder", use_container_width=True, key="open-feedback-folder"):
                    try:
                        open_folder_in_file_manager(FEEDBACK_ROOT / "reports")
                        st.success("Rapportmappen öppnades." if lang == "sv" else "The feedback folder was opened.")
                    except Exception as exc:
                        st.error(f"Kunde inte öppna mappen: {exc}" if lang == "sv" else f"Could not open the folder: {exc}")
            with helper_col_b:
                st.caption(
                    "Mapp: session_data/feedback/reports"
                    if lang == "sv"
                    else "Folder: session_data/feedback/reports"
                )
            with st.form("issue-report-form", clear_on_submit=False):
                st.text_input("Ditt namn" if lang == "sv" else "Your name", key="issue_reporter_name", help="Valfritt." if lang == "sv" else "Optional.")
                issue_col_a, issue_col_b = st.columns(2)
                with issue_col_a:
                    st.selectbox(
                        "Hur stort problem är det?" if lang == "sv" else "How serious is it?",
                        options=["Low", "Medium", "High", "Blocker"],
                        key="issue_severity",
                    )
                with issue_col_b:
                    st.selectbox(
                        "Var hände det?" if lang == "sv" else "Where did it happen?",
                        options=["General", "Setup", "Documents", "Chat", "Summary", "Writing", "Analysis", "Export", "Session save/load", "Performance"],
                        key="issue_area",
                    )
                st.text_input(
                    "Kort namn på problemet" if lang == "sv" else "Short name for the problem",
                    key="issue_title",
                    help="Exempel: Chatten gav fel namn eller PDF gick inte att läsa." if lang == "sv" else "Example: Chat gave the wrong name or PDF would not load.",
                )
                st.text_area(
                    "Vad gick fel?" if lang == "sv" else "What went wrong?",
                    key="issue_what_happened",
                    height=120,
                    help="Beskriv med egna ord." if lang == "sv" else "Describe it in your own words.",
                )
                st.text_area(
                    "Vad hade varit rätt?" if lang == "sv" else "What should have happened?",
                    key="issue_expected",
                    height=90,
                )
                st.text_area(
                    "Vad gjorde du precis innan?" if lang == "sv" else "What did you do right before it happened?",
                    key="issue_steps",
                    height=100,
                    help="Skriv gärna steg för steg om du minns." if lang == "sv" else "Step by step is helpful if you remember.",
                )
                st.text_area(
                    "Vad försökte du få gjort?" if lang == "sv" else "What were you trying to get done?",
                    key="issue_work_context",
                    height=90,
                    help="Till exempel sammanfatta en intervju eller hitta citat." if lang == "sv" else "For example: summarize an interview or find quotes.",
                )
                submitted = st.form_submit_button("Skicka och spara rapport" if lang == "sv" else "Save this report", use_container_width=True)
            if submitted:
                required_fields = [
                    st.session_state.get("issue_title", "").strip(),
                    st.session_state.get("issue_what_happened", "").strip(),
                ]
                if not all(required_fields):
                    st.warning("Fyll i minst problemets namn och vad som gick fel." if lang == "sv" else "Please fill in at least the problem name and what went wrong.")
                else:
                    saved_report = save_tester_issue(build_meta)
                    st.success(
                        (
                            f"Rapport sparad: {saved_report['report_id']}. Den finns i rapportmappen och i listan här nedanför."
                            if lang == "sv"
                            else f"Report saved: {saved_report['report_id']}. It is now in the feedback folder and in the list below."
                        )
                    )
                    st.code(saved_report["report_id"])
                    recent_reports = list_issue_reports(FEEDBACK_ROOT, limit=12)

            report_count = len(recent_reports)
            st.caption(
                f"{report_count} sparade rapporter"
                if lang == "sv"
                else f"{report_count} saved reports"
            )
            if recent_reports:
                bundle_bytes = build_feedback_bundle(FEEDBACK_ROOT, [report["report_id"] for report in recent_reports])
                st.download_button(
                    "Ladda ner rapportpaket" if lang == "sv" else "Download report bundle",
                    data=bundle_bytes,
                    file_name="insikt-feedback-bundle.zip",
                    mime="application/zip",
                    use_container_width=True,
                    key="download-feedback-bundle",
                )
                for report in recent_reports[:6]:
                    label = f"{report.get('created_at', '')[:16]} | {report.get('severity', '')} | {report.get('title', '')}"
                    with st.expander(label, expanded=False):
                        st.caption(report.get("report_id", ""))
                        st.write(report.get("what_happened", ""))
                        st.caption(report.get("markdown_path", ""))
                        markdown_path = Path(report.get("markdown_path", ""))
                        if markdown_path.exists():
                            st.download_button(
                                "Ladda ner rapport" if lang == "sv" else "Download report",
                                data=markdown_path.read_bytes(),
                                file_name=markdown_path.name,
                                mime="text/markdown",
                                key=f"download-report-{report.get('report_id', '')}",
                                use_container_width=True,
                            )

    if st.session_state.processing:
        st.markdown("""
        <div class="processing-overlay">
            <div class="processing-box">
                <h3>""" + get_text("processing") + """</h3>
                <div class="stProgress"><div style="width:100%; background:#d97706; height:4px;"></div></div>
                <p>Bearbetning pågår. Stäng inte appen medan dokumenten förbereds.</p>
            </div>
        </div>
        """, unsafe_allow_html=True)

    lang = st.session_state.lang
    col_left, col_right = st.columns([3, 2])
    with col_left:
        st.markdown(f"""
            <div class="hero">
                <div class="hero-title">{get_text('title')}</div>
                <div class="hero-subtitle">{get_text('hero_subtitle')}</div>
                <div style="margin-top: 0.8rem;">
                    <span class="chip">{get_text('hero_badge_local')}</span>
                    <span class="chip">{get_text('hero_badge_private')}</span>
                    <span class="chip">{get_text('hero_badge_bilingual')}</span>
                </div>
            </div>
            """, unsafe_allow_html=True)
    with col_right:
        st.markdown(f"""
            <div class="card">
                <strong>{'Snabbstart' if lang == 'sv' else 'Quick Start'}</strong>
                <div style="margin-top: 0.6rem; color: #475569;">
                    1) {'Ladda upp filer och bygg kunskapsbasen' if lang == 'sv' else 'Upload files and build the knowledge base'}<br>
                    2) {'Chatta, sammanfatta eller skriv manus' if lang == 'sv' else 'Chat, summarize, or draft scripts'}<br>
                    3) {'Exportera eller fortsätt granska' if lang == 'sv' else 'Export or keep investigating'}
                </div>
            </div>
            """, unsafe_allow_html=True)

    main_col_left, main_col_right = st.columns([5, 4], vertical_alignment="top")

    with main_col_right:
        preview_source = st.session_state.get("selected_preview_source", "")
        preview_page = st.session_state.get("selected_preview_page", "")
        preview_excerpt = st.session_state.get("preview_excerpt", "")
        preview_docs = [doc for doc in (st.session_state.get("raw_pages") or []) if doc.metadata.get("source") == preview_source]
        if st.session_state.get("available_sources"):
            st.markdown("### " + ("Dokumentnavigator" if lang == "sv" else "Document Navigator"))
            navigator_source = st.selectbox(
                "Fil" if lang == "sv" else "File",
                options=st.session_state.get("available_sources", []),
                index=st.session_state.get("available_sources", []).index(preview_source) if preview_source in st.session_state.get("available_sources", []) else 0,
                key="navigator_source",
            )
            if navigator_source != preview_source:
                select_preview_target(navigator_source, origin="manual")
                st.rerun()
            navigator_docs = [doc for doc in (st.session_state.get("raw_pages") or []) if doc.metadata.get("source") == navigator_source]
            st.caption(
                (
                    f"{len(navigator_docs)} sidor tillgängliga"
                    if lang == "sv"
                    else f"{len(navigator_docs)} pages available"
                )
            )
            if st.session_state.get("recent_citations"):
                with st.expander("Senaste källhopp" if lang == "sv" else "Recent citation jumps", expanded=False):
                    for idx, item in enumerate(st.session_state.get("recent_citations", [])):
                        label = f"{item.get('source')} - {('sida' if lang == 'sv' else 'page')} {item.get('page') or '?'}"
                        if st.button(label, key=f"recent-citation-{idx}", use_container_width=True):
                            select_preview_target(item.get("source", ""), item.get("page", ""), excerpt=item.get("excerpt", ""), origin="citation")
                            st.rerun()
            if st.session_state.get("preview_bookmarks"):
                with st.expander("Bokmärken" if lang == "sv" else "Bookmarks", expanded=False):
                    for idx, item in enumerate(st.session_state.get("preview_bookmarks", [])):
                        label = f"{item.get('source')} - {('sida' if lang == 'sv' else 'page')} {item.get('page') or '?'}"
                        col_open, col_remove = st.columns([4, 1])
                        with col_open:
                            if st.button(label, key=f"bookmark-open-{idx}", use_container_width=True):
                                select_preview_target(item.get("source", ""), item.get("page", ""), excerpt=item.get("excerpt", ""), origin="manual")
                                st.rerun()
                        with col_remove:
                            if st.button("X", key=f"bookmark-remove-{idx}", use_container_width=True):
                                bookmarks = st.session_state.get("preview_bookmarks", [])
                                st.session_state.preview_bookmarks = [b for b in bookmarks if not (b.get("source") == item.get("source") and b.get("page") == item.get("page"))]
                                st.rerun()
        if preview_source and preview_docs:
            st.markdown("### " + ("Dokumentförhandsvisning" if lang == "sv" else "Document Preview"))
            page_options = [str(doc.metadata.get("page", "?")) for doc in preview_docs]
            selected_page = st.selectbox("Sida" if lang == "sv" else "Page", options=page_options, index=page_options.index(preview_page) if preview_page in page_options else 0)
            if selected_page != preview_page:
                st.session_state.selected_preview_page = selected_page
                st.session_state.preview_excerpt = ""
                st.session_state.preview_origin = "manual"
            selected_doc = next((doc for doc in preview_docs if str(doc.metadata.get("page", "?")) == selected_page), preview_docs[0])
            st.caption(preview_source)
            bookmark_col, info_col = st.columns([1, 2])
            with bookmark_col:
                if st.button("Bokmärk sida" if lang == "sv" else "Bookmark page", use_container_width=True):
                    add_preview_bookmark(preview_source, selected_page, excerpt=preview_excerpt or selected_doc.page_content[:500])
                    st.rerun()
            with info_col:
                st.caption("Klicka på en källhänvisning för att hoppa hit." if lang == "sv" else "Click a citation to jump here.")
            if st.session_state.get("preview_origin") in {"citation", "register", "snippet"}:
                st.info(
                    "Källan du öppnade visas här till höger. Den markerade passagen visas under filnamnet."
                    if lang == "sv"
                    else "The source you opened appears here on the right. The highlighted passage is shown below the file name."
                )
            if preview_excerpt:
                st.markdown("**Markerad passage**" if lang == "sv" else "**Highlighted passage**")
                st.markdown(
                    f'<div class="source-box" style="border-left: 4px solid #d97706; background: #fff7ed;">{safe_html_fragment(preview_excerpt)}</div>',
                    unsafe_allow_html=True,
                )
            st.text_area("Preview", selected_doc.page_content[:2500], height=220)

    tab_chat, tab_summary, tab_write, tab_analysis, tab_board, tab_export = main_col_left.tabs([
        "Chatt" if lang == "sv" else "Chat",
        "Sammanfatta" if lang == "sv" else "Summary",
        "Skrivstudio" if lang == "sv" else "Writing Studio",
        "Analys" if lang == "sv" else "Analysis",
        "Arbetsyta" if lang == "sv" else "Case board",
        "Exportera" if lang == "sv" else "Export",
    ])

    with tab_chat:
        render_page_header(
            get_text("chat_title"),
            get_text("chat_help"),
            "Om chatt" if lang == "sv" else "About Chat",
            (
                "### Vad den gör\n"
                "- Söker automatiskt i alla uppladdade dokument.\n"
                "- Svarar med källor när underlaget räcker.\n"
                "- Säger tydligt när materialet inte räcker eller måste dubbelkollas.\n"
                "- Kan fortfarande småprata och hjälpa till naturligt utan att hitta på fakta.\n\n"
                "### Bra att veta\n"
                "- Alla uppladdade filer blir automatiskt en del av kunskapsbasen.\n"
                "- Om svaret saknar tydliga källor bör du kontrollera originaldokumentet."
            )
            if lang == "sv"
            else
            (
                "### What it does\n"
                "- Automatically searches across all uploaded documents.\n"
                "- Answers with citations when the material is sufficient.\n"
                "- Clearly says when the material is insufficient or needs review.\n"
                "- Still chats naturally without inventing facts.\n\n"
                "### Good to know\n"
                "- Every uploaded file becomes part of the knowledge base automatically.\n"
                "- If the answer lacks clear citations, check the original document."
            ),
        )
        st.markdown('<div class="chat-shell">', unsafe_allow_html=True)
        for idx, msg in enumerate(st.session_state.chat_history):
            if msg["role"] == "user":
                with st.chat_message("user"):
                    st.markdown(msg["content"])
            else:
                body, notes = split_assistant_content(msg["content"])
                body, register = prettify_citations_for_display(body or msg["content"], lang)
                message_sources = msg.get("sources", [])
                message_issues = msg.get("issues", [])
                with st.chat_message("assistant"):
                    confidence = assess_answer_confidence(
                        msg["content"],
                        records_to_docs(message_sources) if message_sources and isinstance(message_sources[0], dict) else (message_sources or []),
                        message_issues,
                        lang,
                    )
                    render_confidence_banner(confidence, lang)
                    st.markdown(body)
                    render_citation_register(register, lang, key_prefix=f"history-{idx}")
                    if notes or message_issues:
                        st.markdown('<div class="chat-notes">', unsafe_allow_html=True)
                        st.markdown("**Att kontrollera**" if lang == "sv" else "**Check these notes**")
                        for note in notes:
                            st.markdown(f"- {note}")
                        if message_issues:
                            st.markdown("- " + ("Vissa delar av svaret behöver dubbelkontroll." if lang == "sv" else "Parts of the answer may need verification."))
                        st.markdown("</div>", unsafe_allow_html=True)
                    render_source_snippets(message_sources, lang, key_prefix=f"history-{idx}")
        if prompt := st.chat_input(get_text("chat_input"), disabled=st.session_state.processing):
            st.session_state.chat_history.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)
            with st.spinner("Tänker..." if lang == "sv" else "Thinking..."):
                llm = load_llm(st.session_state.device_choice)
                history_lc = [HumanMessage(content=m["content"]) if m["role"] == "user" else AIMessage(content=m["content"]) for m in st.session_state.chat_history[:-1]]
                answer, sources, issues = rag_chat_with_docs(prompt, history_lc, st.session_state.vectorstore, llm, st.session_state.lang)
            st.session_state.chat_history.append({"role": "assistant", "content": answer, "sources": docs_to_records(sources), "issues": issues})
            body, notes = split_assistant_content(answer)
            body, register = prettify_citations_for_display(body or answer, lang)
            with st.chat_message("assistant"):
                confidence = assess_answer_confidence(answer, sources, issues, lang)
                render_confidence_banner(confidence, lang)
                st.markdown(body)
                render_citation_register(register, lang, key_prefix="latest")
                if notes or issues:
                    st.markdown('<div class="chat-notes">', unsafe_allow_html=True)
                    st.markdown("**Att kontrollera**" if lang == "sv" else "**Check these notes**")
                    for note in notes:
                        st.markdown(f"- {note}")
                    if issues:
                        st.markdown("- " + ("Vissa delar av svaret behöver dubbelkontroll." if lang == "sv" else "Parts of the answer may need verification."))
                    st.markdown("</div>", unsafe_allow_html=True)
                render_source_snippets(sources, lang, "chat-latest")
        st.markdown("</div>", unsafe_allow_html=True)

    with tab_summary:
        render_page_header(
            get_text("summarize_title"),
            get_text("summary_help"),
            "Om Sammanfatta" if lang == "sv" else "About Summary",
            (
                "### Vad den gör\n"
                "- Skapar en längre sammanfattning av dina dokument.\n"
                "- Låter dig styra fokus, ton och längd.\n"
                "- Kan förfina sluttexten i ett extra steg.\n\n"
                "### Bra att veta\n"
                "- För stora dokument kan första körningen ta tid.\n"
                "- Längre måltext tar längre tid eftersom fler modellsteg och mer text behöver genereras."
            )
            if lang == "sv"
            else
            (
                "### What it does\n"
                "- Creates a longer summary of your documents.\n"
                "- Lets you control focus, tone, and length.\n"
                "- Can refine the final text in an extra step.\n\n"
                "### Good to know\n"
                "- Large document sets may take time on the first run.\n"
                "- Longer targets take more time because the model has to generate and merge more text."
            ),
        )
        template_labels = get_reporter_template_options(lang)
        current_template_key = st.session_state.get("reporter_template", "")
        current_template_label = REPORTER_TEMPLATES[current_template_key]["title"][lang] if current_template_key in REPORTER_TEMPLATES else template_labels[0]
        template_col_a, template_col_b = st.columns([3, 1])
        with template_col_a:
            selected_template_label = st.selectbox(
                "Reporter-mall" if lang == "sv" else "Reporter template",
                options=template_labels,
                index=template_labels.index(current_template_label) if current_template_label in template_labels else 0,
                key="reporter_template_label",
            )
        with template_col_b:
            if st.button("Använd mall" if lang == "sv" else "Apply template", use_container_width=True):
                template_key = template_key_from_label(selected_template_label, lang)
                if template_key:
                    apply_reporter_template(template_key, lang)
                    st.rerun()
        if not st.session_state.docs:
            st.info(get_text("error_no_docs"))
        else:
            col_settings, col_preview = st.columns([2, 1])
            with col_settings:
                focus = st.text_area(get_text("focus"), placeholder="t.ex. korruption, specifik person..." if lang == "sv" else "e.g., corruption, specific person...", height=90, disabled=st.session_state.processing, key="summary_focus")
                style = st.selectbox(get_text("style"), ["neutral", "investigative", "dramatic", "formal"], disabled=st.session_state.processing, key="summary_style")
                use_refine = st.checkbox(get_text("refine_btn"), value=st.session_state.get("summary_use_refine", True), disabled=st.session_state.processing, key="summary_use_refine")
            with col_preview:
                target_pages = st.slider(get_text("target_pages"), 1, 50, st.session_state.get("summary_target_pages", 5), disabled=st.session_state.processing, key="summary_target_pages")
                words_per_page = st.number_input(get_text("density"), 100, 500, st.session_state.get("summary_words_per_page", 300), disabled=st.session_state.processing, key="summary_words_per_page")
                target_words = int(target_pages * words_per_page)
                st.caption(f"{get_text('summary_estimate')}: {target_words} {'ord' if lang=='sv' else 'words'}")
            if st.button(get_text("summarize_btn"), disabled=st.session_state.processing or st.session_state.summary_running, use_container_width=True):
                st.session_state.summary_stage = "initializing"
                st.session_state.summary_stages_log = []
                st.session_state.summary_percentage = 0
                st.session_state.summary_current_batch = 0
                st.session_state.summary_total_batches = 0
                st.session_state.summary_cancel_requested = False
                start_summary(st.session_state.docs, target_pages, focus, style, words_per_page, st.session_state.lang, use_refine)
                st.rerun()
        if st.session_state.summary_running:
            stage = st.session_state.get("summary_stage", "processing")
            percentage = st.session_state.get("summary_percentage", 0)
            current_batch = st.session_state.get("summary_current_batch", 0)
            total_batches = st.session_state.get("summary_total_batches", 0)
            stage_labels = {"idle": "Väntar" if lang == "sv" else "Idle", "initializing": "Initierar" if lang == "sv" else "Initializing", "processing": "Sammanfattar avsnitt" if lang == "sv" else "Processing sections", "combining": "Kombinerar" if lang == "sv" else "Combining", "polishing": "Färdigställer" if lang == "sv" else "Finalizing", "complete": "Klart" if lang == "sv" else "Complete", "error": "Fel" if lang == "sv" else "Error", "cancelled": "Avbrutet" if lang == "sv" else "Cancelled"}
            st.markdown(f'<div class="info-box"><strong>{stage_labels.get(stage, stage)}</strong><span style="float:right;">{percentage:.0f}%</span></div>', unsafe_allow_html=True)
            st.progress(percentage / 100.0)
            if stage == "processing" and total_batches > 0:
                st.caption(f"Batch {current_batch} / {total_batches}")
            with st.expander("Logg" if lang == "sv" else "Log"):
                for entry in st.session_state.get("summary_stages_log", [])[-10:]:
                    st.caption(f"{entry.get('percentage', 0):.0f}% - {entry.get('message', '')}")
            if st.button("Avbryt" if lang == "sv" else "Cancel", key="cancel_summary"):
                if st.session_state.get("summary_thread"):
                    st.session_state.summary_thread.stop()
                st.session_state.summary_running = False
                st.session_state.summary_stage = "cancelled"
                st.rerun()
        elif st.session_state.summary_error:
            error_msg = st.session_state.summary_error
            if error_msg == "cancelled":
                st.markdown(f'<div class="warning-box">{"Åtgärden avbröts av användaren." if lang=="sv" else "Operation cancelled by user."}</div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="warning-box">{"Fel: " if lang=="sv" else "Error: "}{error_msg}</div>', unsafe_allow_html=True)
        elif st.session_state.summary_result:
            st.markdown("### " + ("Sammanfattning" if lang == "sv" else "Summary"))
            summary_confidence = assess_answer_confidence(
                st.session_state.summary_result,
                st.session_state.get("raw_pages") or [],
                ["missing_citations"] if not ("[Source:" in st.session_state.summary_result or "[Källa:" in st.session_state.summary_result) else [],
                lang,
            )
            render_confidence_banner(summary_confidence, lang)
            summary_body, summary_register = prettify_citations_for_display(st.session_state.summary_result, lang)
            st.markdown(summary_body)
            render_citation_register(summary_register, lang, key_prefix="summary")
            st.session_state.last_summary = st.session_state.summary_result
            st.success(get_text("success_summary"))
            st.download_button("Ladda ner sammanfattning" if lang == "sv" else "Download summary", data=st.session_state.summary_result.encode("utf-8"), file_name="insikt-summary.txt", mime="text/plain")

    with tab_write:
        render_page_header(
            get_text("writing_title"),
            get_text("writing_help"),
            "Om Skrivstudio" if lang == "sv" else "About Writing Studio",
            (
                "### Vad den gör\n"
                "- Använder dina dokument som faktabas för att skriva utkast.\n"
                "- Du väljer roll, format, ton och längd.\n"
                "- Kan skapa artikel, manus eller dokumentärupplägg.\n\n"
                "### Bra att veta\n"
                "- Dokumentär-pipeline kör i flera steg: disposition, scenlista och slutligt manus.\n"
                "- Alla uppladdade dokument används automatiskt som underlag."
            )
            if lang == "sv"
            else
            (
                "### What it does\n"
                "- Uses your documents as a factual base for drafting.\n"
                "- Lets you choose role, format, tone, and length.\n"
                "- Can create articles, scripts, or documentary-style drafts.\n\n"
                "### Good to know\n"
                "- The documentary pipeline runs in stages: outline, scene list, and final script.\n"
                "- All uploaded documents are used automatically as source material."
            ),
        )
        if st.session_state.get("reporter_template") in REPORTER_TEMPLATES:
            template_name = REPORTER_TEMPLATES[st.session_state["reporter_template"]]["title"][lang]
            st.caption(
                f"Aktiv mall: {template_name}"
                if lang == "sv"
                else f"Active template: {template_name}"
            )
        brief = st.text_area(get_text("writing_brief"), placeholder=get_text("writing_placeholder"), height=140, key="writing_brief")
        col_a, col_b = st.columns(2)
        with col_a:
            role_key = st.selectbox(get_text("writing_role"), options=list(WRITING_ROLES.keys()), format_func=lambda x: WRITING_ROLES[x]["sv"] if lang == "sv" else WRITING_ROLES[x]["en"], key="writing_role")
            format_key = st.selectbox(get_text("writing_format"), options=list(WRITING_FORMATS.keys()), format_func=lambda x: WRITING_FORMATS[x]["sv"] if lang == "sv" else WRITING_FORMATS[x]["en"], key="writing_format")
        with col_b:
            tone_key = st.selectbox(get_text("writing_tone"), options=list(WRITING_TONES.keys()), format_func=lambda x: WRITING_TONES[x]["sv"] if lang == "sv" else WRITING_TONES[x]["en"], key="writing_tone")
            length_key = st.selectbox(get_text("writing_length"), options=list(WRITING_LENGTHS.keys()), format_func=lambda x: WRITING_LENGTHS[x]["sv"] if lang == "sv" else WRITING_LENGTHS[x]["en"], key="writing_length")
        use_sources = st.checkbox(get_text("writing_use_sources"), value=st.session_state.get("writing_use_sources", True), key="writing_use_sources")
        use_pipeline = st.checkbox(get_text("writing_pipeline"), value=st.session_state.get("writing_pipeline", False), key="writing_pipeline", help=get_text("writing_pipeline_help"))
        if st.button(get_text("writing_generate"), disabled=st.session_state.processing or not brief):
            llm = load_llm(st.session_state.device_choice)
            result, used_sources = rag_generate_writing(brief, WRITING_ROLES[role_key]["sv" if lang == "sv" else "en"], WRITING_FORMATS[format_key]["sv" if lang == "sv" else "en"], WRITING_TONES[tone_key]["sv" if lang == "sv" else "en"], WRITING_LENGTHS[length_key]["words"], lang, st.session_state.vectorstore, llm, use_sources=use_sources, use_pipeline=use_pipeline)
            st.session_state.writing_result = result
            st.session_state.writing_sources = used_sources
        if st.session_state.get("writing_result"):
            st.markdown("### " + get_text("writing_result"))
            writing_sources = st.session_state.get("writing_sources", [])
            writing_confidence = assess_answer_confidence(
                st.session_state.writing_result,
                writing_sources,
                ["missing_citations"] if st.session_state.get("writing_use_sources", True) and writing_sources and not ("[Source:" in st.session_state.writing_result or "[Källa:" in st.session_state.writing_result) else [],
                lang,
            )
            render_confidence_banner(writing_confidence, lang)
            writing_body, writing_register = prettify_citations_for_display(st.session_state.writing_result, lang)
            st.markdown(writing_body)
            render_citation_register(writing_register, lang, key_prefix="writing")
            render_source_snippets(writing_sources, lang, "writing")
            claim_items = extract_claim_check_items(st.session_state.writing_result, lang)
            flagged_items = [item for item in claim_items if item["needs_review"]]
            with st.expander("Claim checker" if lang == "en" else "Påståendekontroll", expanded=bool(flagged_items)):
                if not flagged_items:
                    st.success(
                        "No obvious unsupported draft claims were flagged."
                        if lang == "en"
                        else "Inga uppenbart osäkra påståenden flaggades i utkastet."
                    )
                else:
                    st.warning(
                        (
                            f"{len(flagged_items)} formuleringar bör verifieras innan publicering."
                            if lang == "sv"
                            else f"{len(flagged_items)} draft statements should be verified before publishing."
                        )
                    )
                    for idx, item in enumerate(flagged_items[:15], start=1):
                        st.markdown(f"**{idx}.** {item['sentence']}")
                        reason_labels = []
                        if "missing_citation" in item["reasons"]:
                            reason_labels.append("saknar källhänvisning" if lang == "sv" else "missing citation")
                        if "speculative_language" in item["reasons"]:
                            reason_labels.append("osäker formulering" if lang == "sv" else "speculative wording")
                        if "number_without_citation" in item["reasons"]:
                            reason_labels.append("siffra utan källa" if lang == "sv" else "number without citation")
                        if reason_labels:
                            st.caption(", ".join(reason_labels))

    with tab_analysis:
        render_page_header(
            get_text("analysis_title"),
            "",
            "Om Analys" if lang == "sv" else "About Analysis",
            (
                "### Vad den gör\n"
                "- Hjälper dig hitta namn, organisationer och platser.\n"
                "- Tar fram datum, nyckelord, citat, källjämförelser och enklare sentiment.\n\n"
                "### Bra att veta\n"
                "- Resultaten är till för översikt och researchstöd.\n"
                "- Dubbelkolla alltid viktiga slutsatser i originalkällan."
            )
            if lang == "sv"
            else
            (
                "### What it does\n"
                "- Helps you find names, organizations, and places.\n"
                "- Pulls out dates, keywords, quotes, source comparisons, and basic sentiment.\n\n"
                "### Good to know\n"
                "- Results are for overview and research support.\n"
                "- Always verify important conclusions in the original source."
            ),
        )
        if not st.session_state.docs:
            st.warning("Ladda upp dokument först." if lang == "sv" else "Please upload documents first.")
        else:
            ner_pipeline = load_ner_pipeline(st.session_state.device_choice)
            sent_pipeline = load_sentiment_pipeline(st.session_state.device_choice)
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                if st.button(get_text("ner_extract"), disabled=st.session_state.processing):
                    with st.spinner("Extraherar..." if lang == "sv" else "Extracting..."):
                        entities = extract_entities(" ".join([d.page_content for d in st.session_state.docs[:50]]), ner_pipeline)
                        for key, values in entities.items():
                            st.markdown(f"**{key}**")
                            st.write(", ".join(list(values)[:20]) + ("..." if len(values) > 20 else ""))
            with col2:
                if st.button(get_text("timeline"), disabled=st.session_state.processing):
                    with st.spinner("Skapar tidslinje..." if lang == "sv" else "Building timeline..."):
                        timeline = extract_timeline(st.session_state.docs[:200])
                        if timeline:
                            for item in timeline[:30]:
                                st.markdown(f"**{item['date']}** - {item['context'][:150]}...\n*{item['source']} {'s.' if lang == 'sv' else 'p.'}{item['page']}*")
                        else:
                            st.info("Inga datum hittades." if lang == "sv" else "No dates found.")
            with col3:
                if st.button(get_text("keywords"), disabled=st.session_state.processing):
                    with st.spinner("Extraherar nyckelord..." if lang == "sv" else "Extracting keywords..."):
                        text = st.session_state.last_summary if st.session_state.last_summary else st.session_state.docs[0].page_content
                        st.write("Toppnyckelord:" if lang == "sv" else "Top keywords:", ", ".join(extract_keywords(text, language="sv" if lang == "sv" else "en")))
            with col4:
                if st.button("Hitta citat" if lang == "sv" else "Find quotes", disabled=st.session_state.processing):
                    with st.spinner("Söker citat..." if lang == "sv" else "Finding quotes..."):
                        st.session_state.quote_candidates = extract_quote_candidates(st.session_state.raw_pages[:200])
            comparison_col_a, comparison_col_b = st.columns([4, 1])
            with comparison_col_a:
                st.text_input(
                    "Jämför person, händelse eller påstående" if lang == "sv" else "Compare person, event, or claim",
                    key="comparison_query",
                    placeholder="t.ex. minister budget 2024" if lang == "sv" else "e.g. minister budget 2024",
                )
            with comparison_col_b:
                if st.button("Jämför källor" if lang == "sv" else "Compare sources", disabled=st.session_state.processing or not st.session_state.get("comparison_query", "").strip(), use_container_width=True):
                    comparison_query = st.session_state.get("comparison_query", "")
                    st.session_state.comparison_results = compare_sources(st.session_state.raw_pages[:300], comparison_query)
                    st.session_state.comparison_last_run = comparison_query
            if st.session_state.get("comparison_results"):
                header_col_a, header_col_b = st.columns([3, 1])
                with header_col_a:
                    st.markdown("### " + ("Källjämförelse" if lang == "sv" else "Source comparison"))
                with header_col_b:
                    if st.button("Rensa" if lang == "sv" else "Clear", key="clear-comparison-results", use_container_width=True):
                        st.session_state.comparison_results = []
                        st.session_state.comparison_query = ""
                        st.session_state.comparison_last_run = ""
                        st.rerun()
                for idx, item in enumerate(st.session_state.get("comparison_results", [])):
                    st.markdown(f"**{item['source']}**")
                    st.caption(f"{'Sida' if lang == 'sv' else 'Page'} {item['page']} | {'träffar' if lang == 'sv' else 'matches'} {item['score']}")
                    st.markdown(
                        f'<div class="source-box">{safe_html_fragment(item["excerpt"])}</div>',
                        unsafe_allow_html=True,
                    )
                    if st.button("Öppna i navigatorn" if lang == "sv" else "Open in navigator", key=f"comparison-open-{idx}", use_container_width=True):
                        select_preview_target(item["source"], str(item["page"]), excerpt=item["excerpt"], origin="snippet")
                        st.rerun()
            elif st.session_state.get("comparison_last_run"):
                st.info(
                    (
                        f"Inga tydliga träffar hittades för '{st.session_state.get('comparison_last_run')}'."
                        if lang == "sv"
                        else f"No clear source matches were found for '{st.session_state.get('comparison_last_run')}'."
                    )
                )
            if st.session_state.get("quote_candidates"):
                header_col_a, header_col_b = st.columns([3, 1])
                with header_col_a:
                    st.markdown("### " + ("Citatkandidater" if lang == "sv" else "Quote candidates"))
                with header_col_b:
                    if st.button("Rensa" if lang == "sv" else "Clear", key="clear-quote-candidates", use_container_width=True):
                        st.session_state.quote_candidates = []
                        st.rerun()
                for idx, item in enumerate(st.session_state.get("quote_candidates", [])[:20]):
                    st.markdown(f'> "{item["quote"]}"')
                    st.caption(
                        f'{item["source"]} | {("sida" if lang == "sv" else "page")} {item["page"]}'
                    )
                    st.caption(item["context"][:220] + ("..." if len(item["context"]) > 220 else ""))
                    if st.button("Öppna källa" if lang == "sv" else "Open source", key=f"quote-open-{idx}", use_container_width=True):
                        select_preview_target(item["source"], str(item["page"]), excerpt=item["quote"], origin="snippet")
                        st.rerun()
            if st.session_state.last_summary and st.button(get_text("sentiment"), disabled=st.session_state.processing):
                with st.spinner("Analyserar sentiment..." if lang == "sv" else "Analyzing sentiment..."):
                    result = analyze_sentiment(st.session_state.last_summary, sent_pipeline)
                    st.write(f"Sentiment: **{result['label']}** (konfidens: {result['score']:.2f})" if lang == "sv" else f"Sentiment: **{result['label']}** (confidence: {result['score']:.2f})")

    with tab_board:
        render_page_header(
            "Arbetsyta" if lang == "sv" else "Case board",
            "Samla arbetsnoter, personer, datum, utdrag och vinklar på ett ställe." if lang == "sv" else "Collect working notes, people, dates, excerpts, and angles in one place.",
            "Om arbetsyta" if lang == "sv" else "About Case board",
            (
                "### Vad den gör\n"
                "- Samlar dina viktigaste arbetsnoter.\n"
                "- Sparar personer, datum, vinklar och källutdrag för fallet.\n"
                "- Följer med när du sparar sessionen.\n\n"
                "### Tips\n"
                "- Klistra in namn eller datum rad för rad.\n"
                "- Lägg till markerade utdrag från dokumentförhandsvisningen."
            )
            if lang == "sv"
            else
            (
                "### What it does\n"
                "- Keeps your core reporting notes in one place.\n"
                "- Stores people, dates, angles, and pinned excerpts for the case.\n"
                "- Travels with the session when you save it.\n\n"
                "### Tips\n"
                "- Add one person or date per line.\n"
                "- Pin highlighted excerpts from the document preview."
            ),
        )
        note_col, list_col = st.columns([2, 1])
        with note_col:
            st.text_area(
                "Arbetsnoter" if lang == "sv" else "Working notes",
                key="case_board_notes",
                height=180,
                placeholder="Skriv vad som verkar viktigt, vad som saknas och vad du vill följa upp."
                if lang == "sv"
                else "Write what seems important, what is missing, and what you want to follow up on.",
            )
            preview_source = st.session_state.get("selected_preview_source", "")
            preview_page = st.session_state.get("selected_preview_page", "")
            preview_excerpt = st.session_state.get("preview_excerpt", "")
            if preview_source and (preview_excerpt or preview_page):
                if st.button("Fäst aktuellt utdrag" if lang == "sv" else "Pin current excerpt", use_container_width=True):
                    excerpt_text = preview_excerpt or next(
                        (
                            doc.page_content[:500]
                            for doc in (st.session_state.get("raw_pages") or [])
                            if doc.metadata.get("source") == preview_source and str(doc.metadata.get("page", "")) == str(preview_page)
                        ),
                        "",
                    )
                    add_case_board_excerpt(preview_source, preview_page, excerpt_text)
                    st.rerun()
        with list_col:
            st.text_area("Viktiga personer" if lang == "sv" else "Key people", key="case_board_people", height=120, placeholder="En person per rad" if lang == "sv" else "One person per line")
            st.text_area("Viktiga datum" if lang == "sv" else "Key dates", key="case_board_dates", height=120, placeholder="Ett datum per rad" if lang == "sv" else "One date per line")
            st.text_area("Vinklar / hypoteser" if lang == "sv" else "Angles / hypotheses", key="case_board_angles", height=120, placeholder="En möjlig vinkel per rad" if lang == "sv" else "One possible angle per line")
        st.markdown("### " + ("Fästa utdrag" if lang == "sv" else "Pinned excerpts"))
        pinned_excerpts = st.session_state.get("case_board_excerpts", [])
        if not pinned_excerpts:
            st.info("Inga utdrag har fästs ännu." if lang == "sv" else "No excerpts have been pinned yet.")
        else:
            clear_col_a, clear_col_b = st.columns([3, 1])
            with clear_col_b:
                if st.button("Rensa alla" if lang == "sv" else "Clear all", key="clear-board-excerpts", use_container_width=True):
                    st.session_state.case_board_excerpts = []
                    st.rerun()
            for idx, item in enumerate(pinned_excerpts):
                st.markdown(f"**{item.get('source', 'Unknown')}** | {('sida' if lang == 'sv' else 'page')} {item.get('page', '?')}")
                st.markdown(
                    f'<div class="source-box">{safe_html_fragment(item.get("excerpt", ""))}</div>',
                    unsafe_allow_html=True,
                )
                action_col_a, action_col_b = st.columns(2)
                with action_col_a:
                    if st.button("Öppna i navigatorn" if lang == "sv" else "Open in navigator", key=f"board-open-{idx}", use_container_width=True):
                        select_preview_target(item.get("source", ""), item.get("page", ""), excerpt=item.get("excerpt", ""), origin="snippet")
                        st.rerun()
                with action_col_b:
                    if st.button("Ta bort" if lang == "sv" else "Remove", key=f"board-remove-{idx}", use_container_width=True):
                        current = st.session_state.get("case_board_excerpts", [])
                        st.session_state.case_board_excerpts = [entry for entry in current if entry != item]
                        st.rerun()

    with tab_export:
        render_page_header(
            get_text("export_title"),
            "",
            "Om export" if lang == "sv" else "About Export",
            (
                "### Vad den gör\n"
                "- Låter dig ladda ner sammanfattningen i flera format.\n"
                "- Kan köra en enkel partiskhetsgranskning.\n"
                "- Kan översätta texten till andra språk.\n\n"
                "### Bra att veta\n"
                "- Export bygger på den senaste sammanfattningen du har skapat.\n"
                "- Kontrollera gärna texten innan du delar den vidare."
            )
            if lang == "sv"
            else
            (
                "### What it does\n"
                "- Lets you download the summary in multiple formats.\n"
                "- Can run a simple bias review.\n"
                "- Can translate the text into other languages.\n\n"
                "### Good to know\n"
                "- Export uses the latest summary you have generated.\n"
                "- It is still worth reviewing the text before sharing it."
            ),
        )
        if not st.session_state.last_summary:
            st.info("Generera en sammanfattning först." if lang == "sv" else "Generate a summary first.")
        else:
            text = st.session_state.last_summary
            safer_exports = st.checkbox(
                "Använd säkrare export med slutnoter och källbilaga" if lang == "sv" else "Use safer export with endnotes and source appendix",
                value=True,
                key="export_safe_mode",
            )
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.download_button("TXT", export_text(text, lang, safer_exports), "sammanfattning.txt" if lang == "sv" else "summary.txt", mime="text/plain")
            with col2:
                st.download_button("DOCX", export_docx(text, lang, safer_exports), "sammanfattning.docx" if lang == "sv" else "summary.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with col3:
                st.download_button("PDF", export_pdf(text, lang, safer_exports), "sammanfattning.pdf" if lang == "sv" else "summary.pdf", mime="application/pdf")
            with col4:
                st.download_button("MD", export_markdown(text, lang, safer_exports), "sammanfattning.md" if lang == "sv" else "summary.md", mime="text/markdown")
            if safer_exports:
                st.caption(
                    "Exporten lägger till slutnoter och en källbilaga baserat på upptäckta hänvisningar."
                    if lang == "sv"
                    else "The export adds endnotes and a source appendix based on detected citations."
                )
            if st.button(get_text("bias_check"), disabled=st.session_state.processing):
                with st.spinner("Kontrollerar partiskhet..." if lang == "sv" else "Checking for bias..."):
                    llm = load_llm(st.session_state.device_choice)
                    st.write(bias_check(text, llm, st.session_state.lang))
            target_lang = st.selectbox(get_text("translate"), ["Svenska", "English", "Spanish", "French", "German"] if lang == "sv" else ["Swedish", "English", "Spanish", "French", "German"])
            if st.button("Översätt" if lang == "sv" else "Translate", disabled=st.session_state.processing):
                with st.spinner("Översätter..." if lang == "sv" else "Translating..."):
                    llm = load_llm(st.session_state.device_choice)
                    translated = translate_text(text, target_lang, llm, st.session_state.lang)
                    st.markdown(f"### Översättning ({target_lang})" if lang == "sv" else f"### Translation ({target_lang})")
                    st.write(translated)

    st.divider()
    st.caption("Insikt - 100% lokalt, privat och gratis." if lang == "sv" else "Insikt - 100% local, private, and free.")
    st.caption(build_meta.get("label", ""))
    if st.session_state.summary_running:
        time.sleep(0.3)
        st.rerun()

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        try:
            build_meta = get_build_metadata()
            save_issue_report(
                FEEDBACK_ROOT,
                {
                    "reporter_name": "system",
                    "severity": "Blocker",
                    "area": "Crash",
                    "title": "Critical app crash",
                    "what_happened": str(e),
                    "expected": "App should stay running.",
                    "steps": "Crash happened during startup or runtime.",
                    "work_context": "Automatic crash capture",
                    "app_version_label": build_meta.get("label", ""),
                    "app_context": {
                        "language": st.session_state.get("lang", "sv") if "lang" in st.session_state else "sv",
                    },
                },
            )
        except Exception:
            pass
        st.error("Ett kritiskt fel uppstod. Starta om appen." if st.session_state.get("lang","sv")=="sv" else "A critical error occurred. Please restart the app.")
        st.exception(e)


